from __future__ import annotations

import csv
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DELIVERY = PROJECT_ROOT / "outputs" / "phase2_delivery"
MULTIPAGE = PROJECT_ROOT / "data" / "multipage" / "finance_40"
MULTI_OUTPUT = PROJECT_ROOT / "outputs" / "multipage_colsmol_top3"
OUTPUT = PROJECT_ROOT / "docs" / "VLM-RAG项目二阶段总结报告.docx"


def main() -> None:
    retrieval = _csv(DELIVERY / "five_retrieval_baselines.csv")
    generation = _csv(DELIVERY / "four_smolvlm_qa_schemes.csv")
    failure_summary = _json(DELIVERY / "retrieval_failure_summary.json")
    failure_cases = _json(DELIVERY / "retrieval_failure_cases.json")
    multipage_manifest = _json(MULTIPAGE / "manifest.json")
    multipage_metrics = _json(MULTI_OUTPUT / "metrics.json")
    multipage_predictions = _json(MULTI_OUTPUT / "predictions.json")
    annotations = _json(MULTIPAGE / "annotation_readable.json")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    _styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    run = p.add_run("页面级文档 VLM-RAG\n项目二阶段总结报告")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(26)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("单页基线复核、失败分析与连续多页联合推理验证")
    run.font.size = Pt(15)
    run.font.name = "KaiTi"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(230)
    p.add_run("实验日期：2026年8月\n数据规模：单页测试500条；多页验证40条").font.size = Pt(12)
    doc.add_page_break()

    _heading(doc, "摘  要", 1)
    _body(doc, (
        "本阶段围绕页面级文档视觉检索增强生成系统开展统一复核与跨页能力验证。首先，将 SigLIP "
        "零样本检索、InfoNCE 部分微调、难负例微调、PPOCR+BGE 与 ColSmol 五种检索方案置于同一 "
        "ChartQA 测试集和同一指标口径下比较；其次，对500条测试 Query 的逐条结果进行失败类型归因；"
        "再次，将 Oracle Page、Top-1、Top-3逐页融合和Top-3拼图四种 SmolVLM 问答方案统一分析。"
        "最后，新建10份连续三页财务报告和40条必须联合两个证据页才能回答的问题，并在本地 RTX 4070 "
        "上完成 ColSmol Top-3 检索和 SmolVLM-500M 多图问答验证。结果显示，ColSmol 在单页测试上以 "
        "Recall@3=0.3900、MRR@10=0.3669 获得最佳检索性能；多页验证中任意证据召回率为1.0000，全部"
        "证据召回率为0.9750，但多图联合问答宽松准确率为0.0000。这一结果表明：当前系统已具备较强的"
        "页面定位能力，但现有轻量生成模型不能可靠完成跨页取值、对齐与算术推理。"
    ))
    _body(doc, "关键词：视觉RAG；页面检索；ColSmol；InfoNCE；SmolVLM；多页问答；证据覆盖")

    _heading(doc, "1 研究任务与阶段边界", 1)
    _body(doc, (
        "此前实验以 ChartQA 单张图表页面为基本样本。Top-3 仅表示从整个页面库中返回三个候选页面，"
        "并不等于一份真实文档中的连续多页联合推理。本阶段因此分成两部分：一是整理并复核已有单页"
        "实验；二是增加具有明确文档归属、连续页码和多证据标注的跨页验证。所有既有指标均来自实际"
        "实验产物；新多页数据为程序化构造的受控验证集，不冒充真实企业数据。"
    ))

    _heading(doc, "2 五种检索方案统一基线", 1)
    _body(doc, (
        "五种方案使用相同的 ChartQA 大规模子集，其中训练集5000条、开发集500条、测试集500条，"
        "页面库包含2958张唯一图表页面。Recall@K 表示前K个结果中是否包含正确证据页，MRR@10同时"
        "考虑正确页是否出现以及出现位置。"
    ))
    _table(doc, retrieval, ["方案", "作用", "Recall@1", "Recall@3", "Recall@10", "MRR@10"], [3.6, 6.2, 2, 2, 2, 2])
    _body(doc, (
        "ColSmol 的 Recall@1 为0.3320，相比 SigLIP 零样本的0.1600提高0.1720；MRR@10由0.1903"
        "提高到0.3669。InfoNCE 部分微调使 SigLIP 的 Recall@10 从0.2520提高到0.3240，证明领域配对"
        "训练有效。加入一个难负例后没有继续提升，MRR@10略降至0.2313，说明难负例质量和采样强度"
        "需要谨慎控制。OCR+BGE 的 Recall@1 达0.1980，但 Recall@10仅0.2800，整体弱于ColSmol。"
    ))

    _heading(doc, "3 检索失败案例原因分析", 1)
    counts = failure_summary["category_counts"]
    failure_table = [
        {"类别": key, "数量": value, "占比": f"{value / failure_summary['test_query_count']:.1%}", "解释": failure_summary["interpretation"][key]}
        for key, value in counts.items()
    ]
    _table(doc, failure_table, ["类别", "数量", "占比", "解释"], [3.4, 1.7, 1.8, 10.5])
    _body(doc, (
        "最大问题是254条样本在四条具有逐条结果的检索链路中均未进入前10，说明仅靠扩大Top-K无法"
        "解决这些样本。54条样本由ColSmol单独救回前三，体现局部视觉词元晚交互的优势；30条样本由"
        "InfoNCE微调把证据页提升到前三，说明微调主要改善部分边界样本；另有3条OCR路径排名更靠前，"
        "提示融合OCR精确文本匹配仍有价值。难负例方案没有保存逐Query结果，因此失败归因采用其余四条"
        "可追溯链路，难负例仅进入总体基线表。"
    ))
    for case in failure_cases[:8]:
        _heading(doc, f"案例 {case['query_id']}：{case['failure_category']}", 2)
        _body(doc, (
            f"Query：{case['query']}。证据页排名分别为 SigLIP零样本 {case['siglip_zero_rank']}、"
            f"微调SigLIP {case['siglip_tuned_rank']}、OCR+BGE {case['ocr_bge_rank']}、"
            f"ColSmol {case['colsmol_rank']}。原因判断：{case['reason']}。"
        ))

    _heading(doc, "4 四种 SmolVLM 问答方案比较", 1)
    _table(doc, generation, ["方案", "作用", "证据召回率", "完全匹配率", "宽松准确率", "命中证据后的准确率", "平均生成耗时(ms)"], [3, 4.6, 2, 2, 2, 2.7, 2.7])
    _body(doc, (
        "Oracle Page 的宽松准确率只有0.3460，说明即使证据完全正确，SmolVLM-500M仍是明显瓶颈。"
        "ColSmol Top-1 的宽松准确率为0.1840，Top-3逐页融合仅升至0.1860，却把平均耗时从455.1 ms"
        "提高到1369.5 ms。Top-3拼图虽只需350.3 ms，但准确率降到0.1260，表明缩放和视觉拥挤造成"
        "信息损失。现有单页测试下，ColSmol Top-1是较合理的速度—准确率折中。"
    ))

    _heading(doc, "5 连续多页小型验证集", 1)
    _body(doc, (
        f"新验证集包含{multipage_manifest['document_count']}份财务报告、{multipage_manifest['page_count']}张"
        f"连续页面和{multipage_manifest['sample_count']}条问答。每份报告固定三页：第1页记录2025年收入、"
        "成本和员工数；第2页记录2026年对应数据；第3页记录研发支出等补充指标。四类问题分别考察收入"
        "增量、经营利润变化、员工增量，以及2026年经营利润与研发支出的跨页求和。每题标注两个必要"
        "证据页，并设置 requires_all_evidence_pages=true。"
    ))
    example_pages = [MULTIPAGE / "pages" / f"finance_01_aster_p{i}.png" for i in (1, 2, 3)]
    for index, image_path in enumerate(example_pages, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Cm(8.2))
        cap = doc.add_paragraph(f"图5-{index}  Aster Holdings 连续财务报告第{index}页")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = doc.styles["Caption"]
    examples = annotations[:4]
    _table(doc, examples, ["query_id", "query", "answer", "evidence_pages", "derivation"], [3.2, 7.2, 1.5, 4, 5])

    _heading(doc, "6 多页检索与联合问答结果", 1)
    multi_table = [
        {"指标": "任意证据召回率", "数值": f"{multipage_metrics['any_evidence_recall']:.4f}", "含义": "Top-3至少包含一个必要证据页"},
        {"指标": "全部证据召回率", "数值": f"{multipage_metrics['all_evidence_recall']:.4f}", "含义": "Top-3同时包含全部必要证据页"},
        {"指标": "平均证据覆盖率", "数值": f"{multipage_metrics['mean_evidence_coverage']:.4f}", "含义": "必要证据页被召回的平均比例"},
        {"指标": "宽松问答准确率", "数值": f"{multipage_metrics['relaxed_accuracy']:.4f}", "含义": "SmolVLM多图联合回答的数值正确率"},
        {"指标": "平均生成耗时", "数值": f"{multipage_metrics['mean_generation_ms']:.1f} ms", "含义": "RTX 4070上每条多页问题平均耗时"},
    ]
    _table(doc, multi_table, ["指标", "数值", "含义"], [5, 3, 10])
    _body(doc, (
        "ColSmol在受控数据上可以较稳定地定位跨页证据：40题全部命中至少一个证据页，39题在Top-3中"
        "同时找齐两个必要证据页。但SmolVLM-500M的40题宽松准确率为0。逐条输出显示，它经常直接抄写"
        "2025年或2026年的单页值，而没有执行减法；有时把员工总数当成员工增量；也会输出算式但使用"
        "错误操作数。即便仅统计39条已找齐全部证据的样本，准确率仍为0，因而失败主要发生在跨页对齐"
        "与计算阶段，而不是页面检索阶段。"
    ))
    by_type = _multipage_type_summary(multipage_predictions)
    _table(doc, by_type, ["问题类型", "数量", "正确数", "准确率", "典型错误"], [4, 2, 2, 2, 8])

    _heading(doc, "7 本阶段完成的工程工作", 1)
    _body(doc, (
        "本阶段新增可重复生成的连续多页数据构建脚本，形成 pages.jsonl、samples.jsonl、manifest.json 和"
        "可读标注文件；新增真正面向多证据问题的评测脚本，将任意证据命中、全部证据命中和证据覆盖率"
        "分开统计，并支持逐条落盘和断点续跑；新增实验汇总脚本，可从原始JSON自动生成五种检索基线、"
        "四种问答方案和失败案例分析。所有多页页面、Query、答案、证据页及推导规则均可追溯。"
    ))

    _heading(doc, "8 阶段结论", 1)
    _body(doc, (
        "本阶段完成了从单页实验总结到连续多页验证的实质性推进。单页结果确认ColSmol是当前最强检索"
        "方案，InfoNCE微调对SigLIP有稳定增益，但简单加入难负例未带来进一步提升。多页结果进一步说明，"
        "‘Top-K候选页问答’与‘真正跨页联合推理’不能混为一谈：前者只要求候选中存在答案页，后者要求"
        "同时找齐多个证据页并正确完成数值对齐与运算。当前系统在受控验证集上已基本跨过多证据检索门槛，"
        "但SmolVLM-500M未跨过联合推理门槛。因此现阶段可以宣称完成页面级视觉RAG原型和多证据检索验证，"
        "不能宣称已经解决真实多页财务问答。"
    ))

    _page_numbers(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


def _multipage_type_summary(rows: list[dict[str, object]]) -> list[dict[str, object]]:
    groups: dict[str, list[dict[str, object]]] = defaultdict(list)
    for row in rows:
        query_id = row["query_id"]
        number = int(query_id.rsplit("q", 1)[1])
        label = {1: "收入增量", 2: "经营利润变化", 3: "员工增量", 4: "利润与研发求和"}[number]
        groups[label].append(row)
    errors = {
        "收入增量": "抄写单年收入，未执行跨页减法",
        "经营利润变化": "未分别计算两年利润或错误相减",
        "员工增量": "输出员工总数而非年度增量",
        "利润与研发求和": "混淆收入、利润和研发支出",
    }
    return [
        {
            "问题类型": label,
            "数量": len(values),
            "正确数": sum(bool(row["relaxed_correct"]) for row in values),
            "准确率": f"{sum(bool(row['relaxed_correct']) for row in values) / len(values):.1%}",
            "典型错误": errors[label],
        }
        for label, values in groups.items()
    ]


def _styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Pt(22)
    for name, size in (("Title", 26), ("Heading 1", 18), ("Heading 2", 14)):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
    doc.styles["Caption"].font.name = "KaiTi"
    doc.styles["Caption"]._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    doc.styles["Caption"].font.size = Pt(10.5)


def _heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(8)


def _body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _table(doc: Document, rows: list[dict[str, object]], columns: list[str], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, column in enumerate(columns):
        cell = table.rows[0].cells[index]
        cell.text = column
        cell.width = Cm(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            value = row.get(column, "")
            if isinstance(value, list):
                value = ", ".join(str(item) for item in value)
            cells[index].text = str(value)
            cells[index].width = Cm(widths[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if index != 1 else WD_ALIGN_PARAGRAPH.LEFT
            for run in cells[index].paragraphs[0].runs:
                run.font.size = Pt(9)
    _three_line(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _three_line(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "insideH", "bottom"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "12" if edge in ("top", "bottom") else "4")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    for edge in ("left", "right", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def _page_numbers(doc: Document) -> None:
    for section in doc.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend((begin, instr, end))


def _csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        return list(csv.DictReader(file))


def _json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
