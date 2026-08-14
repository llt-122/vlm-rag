from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyArrowPatch, Rectangle
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from generate_final_project_report import (
    add_caption,
    add_figure,
    add_page_field,
    add_toc,
    configure_document,
    set_cell_margins,
    set_repeat_table_header,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "docs" / "report_assets"
CASE_DIR = ASSET_DIR / "real_cases"
OUTPUT = ROOT / "docs" / "VLM-RAG项目报告_排版放大修订版.docx"
TITLE = "基于视觉语言模型的页面级文档检索与问答实验研究"


def paragraph(document: Document, text: str, *, font: str = "宋体", size: float = 12,
              first_line: bool = True, bold_prefix: str | None = None,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, before: float = 0, after: float = 0) -> None:
    p = document.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.85)
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, east_asia="黑体", size=size, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, east_asia=font, size=size)
    else:
        run = p.add_run(text)
        set_run_font(run, east_asia=font, size=size)


def analysis_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.right_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "9E9E9E")
    borders.append(bottom)
    p_pr.append(borders)
    lead = p.add_run("实验分析：")
    set_run_font(lead, east_asia="楷体", size=12, bold=True, color="333333")
    run = p.add_run(text)
    set_run_font(run, east_asia="楷体", size=12, color="333333")


def code_module(document: Document, files: str, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.right_indent = Cm(0.25)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    lead = p.add_run(files)
    set_run_font(lead, east_asia="楷体", western="Consolas", size=11.3, bold=True, color="314A5E")
    body = p.add_run("。" + text)
    set_run_font(body, east_asia="宋体", size=11.5)


def case_info(document: Document, query: str, answer: str, evidence: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.right_indent = Cm(0.7)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shd)
    for label, value in (("真实 Query：", query), ("标准答案：", answer), ("证据页：", evidence)):
        lead = p.add_run(label)
        set_run_font(lead, east_asia="楷体", size=11.5, bold=True)
        run = p.add_run(value + ("\n" if label != "证据页：" else ""))
        set_run_font(run, east_asia="楷体", size=11.5)


def equation(document: Document, value: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(value)
    set_run_font(run, east_asia="宋体", western="Cambria Math", size=11.5)


def add_heading(document: Document, text: str, level: int) -> None:
    p = document.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt({1: 15, 2: 10, 3: 7}[level])
    p.paragraph_format.space_after = Pt({1: 8, 2: 5, 3: 4}[level])
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", size={1: 18, 2: 15, 3: 13}[level], bold=True)


def configure_report_styles(document: Document) -> None:
    for level, size in ((1, 18), (2, 15), (3, 13)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    for level, size in ((1, 13), (2, 12), (3, 11)):
        name = f"TOC {level}"
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = level == 1
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.paragraph_format.space_before = Pt(1.5)
        style.paragraph_format.space_after = Pt(1.5)
        style.paragraph_format.line_spacing = 1.15


def add_page_numbers_only(document: Document) -> None:
    for section in document.sections:
        for header in (section.header, section.first_page_header):
            for p in header.paragraphs:
                p.clear()
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.clear()
        add_page_field(fp)
        first_footer = section.first_page_footer
        for p in first_footer.paragraphs:
            p.clear()


def _border(parent, edge: str, *, value: str, size: int = 0, color: str = "000000") -> None:
    node = OxmlElement(f"w:{edge}")
    node.set(qn("w:val"), value)
    if value != "nil":
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
    parent.append(node)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None,
              font_size: float = 10.5) -> None:
    """Add a journal-style three-line table without shading or vertical rules."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    old_borders = tbl_pr.find(qn("w:tblBorders"))
    if old_borders is not None:
        tbl_pr.remove(old_borders)
    borders = OxmlElement("w:tblBorders")
    _border(borders, "top", value="single", size=12)
    _border(borders, "left", value="nil")
    _border(borders, "bottom", value="single", size=12)
    _border(borders, "right", value="nil")
    _border(borders, "insideH", value="nil")
    _border(borders, "insideV", value="nil")
    tbl_pr.append(borders)

    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=90, start=90, bottom=90, end=90)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        _border(tc_borders, "bottom", value="single", size=8)
        tc_pr.append(tc_borders)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(value)
        set_run_font(run, east_asia="宋体", size=font_size, bold=True)
        if widths:
            cell.width = Cm(widths[idx])

    for row_values in rows:
        row = table.add_row()
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=75, start=90, bottom=75, end=90)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(value)
            set_run_font(run, east_asia="宋体", size=font_size)
            if widths:
                cell.width = Cm(widths[idx])

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.35


def page_break(document: Document) -> None:
    document.add_page_break()


def setup_plot_font() -> None:
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    font_manager.fontManager.addfont(str(font_path))
    plt.rcParams["font.family"] = font_manager.FontProperties(fname=str(font_path)).get_name()
    plt.rcParams["axes.unicode_minus"] = False


def make_research_flow(path: Path) -> None:
    setup_plot_font()
    fig, ax = plt.subplots(figsize=(10.0, 6.0))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    labels = [
        "ChartQA 数据\n页面 + Query + 答案", "三条基线\nSigLIP / OCR+BGE / ColSmol",
        "检索器训练\n部分微调 + InfoNCE", "消融实验\n困难负例 0/1/2",
        "端到端问答\nTop-1 / Top-3 / 拼接", "指标与案例\n检索、生成、时延、错误",
    ]
    positions = [(0.35, 3.35), (3.68, 3.35), (7.01, 3.35), (7.01, 1.05), (3.68, 1.05), (0.35, 1.05)]
    for idx, ((x, y), label) in enumerate(zip(positions, labels)):
        ax.add_patch(Rectangle((x, y), 2.55, 1.35, facecolor="white", edgecolor="#333333", linewidth=1.6))
        ax.text(x + 1.275, y + 0.675, label, ha="center", va="center", fontsize=16, linespacing=1.35)
        if idx < len(positions) - 1:
            nx, ny = positions[idx + 1]
            if idx in (0, 1):
                start, end = (x + 2.55, y + 0.675), (nx - 0.10, ny + 0.675)
            elif idx == 2:
                start, end = (x + 1.275, y), (nx + 1.275, ny + 1.45)
            else:
                start, end = (x, y + 0.675), (nx + 2.65, ny + 0.675)
            ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=18, color="#555555", lw=1.5))
    ax.text(0.35, 5.45, "研究主线：建立对照 → 训练检索器 → 验证问答收益", fontsize=20, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_baseline_figure(path: Path) -> None:
    setup_plot_font()
    methods = ["SigLIP\n零样本", "OCR+BGE", "SigLIP\n部分微调", "ColSmol"]
    r1 = [0.160, 0.198, 0.190, 0.332]
    r3 = [0.216, 0.254, 0.260, 0.390]
    r10 = [0.252, 0.280, 0.324, 0.446]
    mrr = [0.1903, 0.2274, 0.2320, 0.3669]
    x = range(4)
    w = 0.18
    colors = ["#D9D9D9", "#A6A6A6", "#6F8799", "#314A5E"]
    fig, ax = plt.subplots(figsize=(9.4, 4.8))
    for j, (name, values) in enumerate((("Recall@1", r1), ("Recall@3", r3), ("Recall@10", r10), ("MRR@10", mrr))):
        bars = ax.bar([i + (j - 1.5) * w for i in x], values, w, label=name, color=colors[j], edgecolor="#333333", linewidth=0.35)
        ax.bar_label(bars, fmt="%.3f", fontsize=7.7, padding=2)
    ax.set_xticks(list(x), methods)
    ax.set_ylim(0, 0.50)
    ax.set_ylabel("指标值")
    ax.set_title("测试集检索结果：三条基线与微调方法")
    ax.grid(axis="y", linestyle="--", linewidth=0.5, alpha=0.35)
    ax.legend(ncol=4, frameon=False, loc="upper left")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_training_figure(path: Path) -> None:
    setup_plot_font()
    epochs = [1, 2, 3, 4]
    loss = [2.072257, 1.618936, 1.435246, 1.368994]
    mrr = [0.251079, 0.255093, 0.259061, 0.258815]
    fig, ax1 = plt.subplots(figsize=(8.8, 4.4))
    ax2 = ax1.twinx()
    a = ax1.plot(epochs, loss, marker="o", color="#314A5E", lw=1.8, label="训练损失")
    b = ax2.plot(epochs, mrr, marker="s", color="#8A4F3D", lw=1.8, linestyle="--", label="验证集 MRR@10")
    ax1.set_xlabel("Epoch")
    ax1.set_ylabel("平均训练损失")
    ax2.set_ylabel("验证集 MRR@10")
    ax1.set_xticks(epochs)
    ax1.set_ylim(1.25, 2.2)
    ax2.set_ylim(0.245, 0.263)
    ax1.grid(axis="y", linestyle="--", linewidth=0.5, alpha=0.35)
    ax1.legend(a + b, [x.get_label() for x in a + b], frameon=False, loc="center right")
    ax1.set_title("SigLIP 部分微调训练曲线")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_generation_figure(path: Path) -> None:
    setup_plot_font()
    labels = ["Oracle", "SigLIP\nTop-1", "SigLIP\nTop-3逐页", "ColSmol\nTop-1", "ColSmol\nTop-3逐页", "ColSmol\nTop-3拼接"]
    acc = [0.346, 0.162, 0.164, 0.184, 0.186, 0.126]
    latency = [466.3, 465.0, 1387.0, 455.1, 1369.5, 350.3]
    fig, ax1 = plt.subplots(figsize=(10.0, 4.8))
    bars = ax1.bar(range(6), acc, color=["#BFBFBF", "#9FB3C1", "#6F8799", "#8FA68E", "#617B62", "#C5B38E"], edgecolor="#333333", linewidth=0.4)
    ax1.bar_label(bars, fmt="%.3f", fontsize=8.2, padding=2)
    ax1.set_xticks(range(6), labels)
    ax1.set_ylabel("宽松准确率")
    ax1.set_ylim(0, 0.39)
    ax1.grid(axis="y", linestyle="--", linewidth=0.5, alpha=0.35)
    ax2 = ax1.twinx()
    ax2.plot(range(6), latency, color="#7A3232", marker="o", lw=1.7, label="平均时延")
    ax2.set_ylabel("平均生成时延（ms）")
    ax2.set_ylim(0, 1600)
    ax1.set_title("端到端问答准确率与时延")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_dataset_figure(path: Path) -> None:
    setup_plot_font()
    fig, ax = plt.subplots(figsize=(9.0, 5.8))
    ax.set_xlim(0, 9.0)
    ax.set_ylim(0, 5.8)
    ax.axis("off")
    ax.text(0.25, 5.35, "ChartQA 实验数据的整理结果", fontsize=20, fontweight="bold")
    cards = [
        (0.25, 3.60, "2,958", "去重图表页面"),
        (3.20, 3.60, "6,000", "真实 Query—答案对"),
        (6.15, 3.60, "5,000", "训练样本"),
        (1.75, 2.05, "500", "验证样本"),
        (4.70, 2.05, "500", "测试样本"),
    ]
    for x, y, value, label in cards:
        ax.add_patch(Rectangle((x, y), 2.55, 1.10, facecolor="white", edgecolor="#3F3F3F", linewidth=1.4))
        ax.text(x + 1.275, y + 0.70, value, ha="center", va="center", fontsize=21, fontweight="bold", color="#314A5E")
        ax.text(x + 1.275, y + 0.28, label, ha="center", va="center", fontsize=15)
    ax.add_patch(Rectangle((0.35, 0.42), 8.30, 0.92, facecolor="#F3F3F3", edgecolor="#666666", linewidth=1.1))
    ax.text(4.50, 0.88, "统一字段：page_id  ·  image_path  ·  query  ·  answers  ·  evidence_page_ids",
            ha="center", va="center", fontsize=14)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_baseline_architecture(path: Path) -> None:
    setup_plot_font()
    fig, ax = plt.subplots(figsize=(9.0, 7.4))
    ax.set_xlim(0, 9.0)
    ax.set_ylim(0, 7.4)
    ax.axis("off")
    ax.text(0.2, 7.00, "三条基线对应三种页面表示方式", fontsize=20, fontweight="bold")
    lanes = [
        (4.95, "SigLIP 零样本\n同架构基础对照", ["Query /\n页面图像", "双塔\n编码", "全局\n单向量", "余弦\nTop-K"]),
        (2.85, "PP-OCR + BGE\n传统文本链路", ["页面\n图像", "PP-OCR\n文本", "BGE\n单向量", "余弦\nTop-K"]),
        (0.75, "ColSmol\n强视觉检索对照", ["Query /\n页面图像", "视觉 token\n编码", "多向量\n表示", "晚交互\nTop-K"]),
    ]
    for y, title, steps in lanes:
        ax.text(0.15, y + 0.55, title, ha="left", va="center", fontsize=14.5, fontweight="bold", linespacing=1.3)
        for idx, step in enumerate(steps):
            x = 2.15 + idx * 1.70
            face = "#F5F5F5" if idx not in (2, 3) else "#E7ECEF"
            ax.add_patch(Rectangle((x, y), 1.35, 1.10, facecolor=face, edgecolor="#444444", linewidth=1.25))
            ax.text(x + 0.675, y + 0.55, step, ha="center", va="center", fontsize=13.5, linespacing=1.25)
            if idx < 3:
                ax.add_patch(FancyArrowPatch((x + 1.35, y + 0.55), (x + 1.66, y + 0.55),
                                             arrowstyle="-|>", mutation_scale=16, color="#666666", lw=1.2))
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_finetune_architecture(path: Path) -> None:
    setup_plot_font()
    fig, ax = plt.subplots(figsize=(9.0, 6.5))
    ax.set_xlim(0, 9.0)
    ax.set_ylim(0, 6.5)
    ax.axis("off")
    ax.text(0.25, 6.05, "SigLIP 部分微调：只更新两塔末端层和投影头", fontsize=20, fontweight="bold")
    for x, title, source in ((0.45, "文本塔", "Query"), (5.75, "视觉塔", "Page image")):
        ax.text(x + 1.4, 5.25, title, ha="center", fontsize=17, fontweight="bold")
        ax.add_patch(Rectangle((x, 4.20), 2.8, 0.78, facecolor="#F7F7F7", edgecolor="#555555", linewidth=1.2))
        ax.text(x + 1.4, 4.59, source, ha="center", va="center", fontsize=16)
        ax.add_patch(Rectangle((x, 2.95), 2.8, 0.82, facecolor="#D9D9D9", edgecolor="#555555", linewidth=1.2))
        ax.text(x + 1.4, 3.36, "冻结的底层\nTransformer", ha="center", va="center", fontsize=15, linespacing=1.15)
        ax.add_patch(Rectangle((x, 1.65), 2.8, 0.86, facecolor="#B7C7D3", edgecolor="#3E5566", linewidth=1.35))
        ax.text(x + 1.4, 2.08, "解冻末端 2 层\n＋投影头", ha="center", va="center", fontsize=15, linespacing=1.15)
        ax.add_patch(FancyArrowPatch((x + 1.4, 4.20), (x + 1.4, 3.80), arrowstyle="-|>", mutation_scale=16, color="#555555"))
        ax.add_patch(FancyArrowPatch((x + 1.4, 2.95), (x + 1.4, 2.54), arrowstyle="-|>", mutation_scale=16, color="#555555"))
    ax.add_patch(FancyArrowPatch((3.25, 2.08), (3.92, 2.08), arrowstyle="-|>", mutation_scale=16, color="#555555"))
    ax.add_patch(FancyArrowPatch((5.75, 2.08), (5.08, 2.08), arrowstyle="-|>", mutation_scale=16, color="#555555"))
    ax.add_patch(Rectangle((3.92, 1.56), 1.16, 1.04, facecolor="#F1E8E2", edgecolor="#765548", linewidth=1.25))
    ax.text(4.50, 2.08, "相似度\n矩阵", ha="center", va="center", fontsize=14.5)
    ax.add_patch(FancyArrowPatch((4.50, 1.55), (4.50, 1.13), arrowstyle="-|>", mutation_scale=16, color="#555555"))
    ax.text(4.50, 0.78, "多正例 InfoNCE（τ=0.07）", ha="center", va="center", fontsize=16, fontweight="bold")
    ax.text(7.10, 0.78, "可训练参数 17.74%", ha="center", va="center", fontsize=15)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_generation_pipeline(path: Path) -> None:
    setup_plot_font()
    fig, ax = plt.subplots(figsize=(9.0, 7.0))
    ax.set_xlim(0, 9.0)
    ax.set_ylim(0, 7.0)
    ax.axis("off")
    ax.text(0.25, 6.55, "检索结果如何进入视觉问答模块", fontsize=20, fontweight="bold")
    top_boxes = [
        (0.35, "Query", "#F5F5F5"),
        (3.42, "检索 Top-K", "#E7ECEF"),
        (6.49, "候选页面", "#F5F5F5"),
    ]
    for idx, (x, label, face) in enumerate(top_boxes):
        ax.add_patch(Rectangle((x, 5.10), 2.15, 0.90, facecolor=face, edgecolor="#444444", linewidth=1.3))
        ax.text(x + 1.075, 5.55, label, ha="center", va="center", fontsize=16)
        if idx < 2:
            ax.add_patch(FancyArrowPatch((x + 2.15, 5.55), (top_boxes[idx + 1][0] - 0.08, 5.55),
                                         arrowstyle="-|>", mutation_scale=17, color="#555555"))
    paths = [
        (3.80, "Top-1\n单页输入", "SmolVLM\n一次生成", "短答案"),
        (2.35, "Top-3\n逐页输入", "3 个候选答案", "分数加权\n融合答案"),
        (0.90, "Top-3\n纵向拼接", "SmolVLM\n一次生成", "短答案"),
    ]
    for y, label, middle, end in paths:
        ax.add_patch(Rectangle((0.35, y), 2.15, 1.00, facecolor="white", edgecolor="#555555", linewidth=1.2))
        ax.text(1.425, y + 0.50, label, ha="center", va="center", fontsize=14.5, linespacing=1.2)
        ax.add_patch(FancyArrowPatch((2.50, y + 0.50), (3.32, y + 0.50), arrowstyle="-|>", mutation_scale=15, color="#666666"))
        ax.add_patch(Rectangle((3.37, y), 2.30, 1.00, facecolor="#F3F3F3", edgecolor="#555555", linewidth=1.2))
        ax.text(4.52, y + 0.50, middle, ha="center", va="center", fontsize=14.5, linespacing=1.2)
        ax.add_patch(FancyArrowPatch((5.67, y + 0.50), (6.49, y + 0.50), arrowstyle="-|>", mutation_scale=15, color="#666666"))
        ax.add_patch(Rectangle((6.54, y), 2.10, 1.00, facecolor="#E7ECEF", edgecolor="#555555", linewidth=1.2))
        ax.text(7.59, y + 0.50, end, ha="center", va="center", fontsize=14.5, linespacing=1.2)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_environment_figure(path: Path) -> None:
    setup_plot_font()
    fig, ax = plt.subplots(figsize=(9.0, 7.0))
    ax.set_xlim(0, 9.0)
    ax.set_ylim(0, 7.0)
    ax.axis("off")
    ax.text(0.25, 6.55, "实际训练与评测环境", fontsize=20, fontweight="bold")
    items = [
        (0.45, 3.85, "基础镜像", "Ubuntu 22.04\nPython 3.12\nPyTorch 2.8.0\nCUDA 12.8"),
        (4.75, 3.85, "阶段一", "RTX 5090 32GB\nSigLIP 基线\nSigLIP 部分微调\n缓存置于数据盘"),
        (4.75, 0.75, "实例迁移", "复制 /root/autodl-tmp\n保留数据与缓存\n保留索引与检查点\n模型代码不变"),
        (0.45, 0.75, "阶段二", "RTX 4090 24GB\nPP-OCR + BGE\nColSmol 与 Stage 4\n统一环境测量时延"),
    ]
    for idx, (x, y, title, detail) in enumerate(items):
        ax.add_patch(Rectangle((x, y), 3.80, 2.20, facecolor="white" if idx % 2 == 0 else "#F3F3F3",
                               edgecolor="#444444", linewidth=1.35))
        ax.text(x + 1.90, y + 1.82, title, ha="center", va="center", fontsize=17, fontweight="bold")
        ax.text(x + 1.90, y + 0.88, detail, ha="center", va="center", fontsize=14.5, linespacing=1.35)
        if idx < len(items) - 1:
            nx, ny = items[idx + 1][0], items[idx + 1][1]
            if idx == 0:
                start, end = (x + 3.80, y + 1.10), (nx - 0.10, ny + 1.10)
            elif idx == 1:
                start, end = (x + 1.90, y), (nx + 1.90, ny + 2.30)
            else:
                start, end = (x, y + 1.10), (nx + 3.90, ny + 1.10)
            ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=17, color="#666666", lw=1.3))
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def figures() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    values = {
        "flow": ASSET_DIR / "research_flow_academic.png",
        "baseline": ASSET_DIR / "baseline_comparison_academic.png",
        "training": ASSET_DIR / "training_curve_academic.png",
        "generation": ASSET_DIR / "generation_tradeoff_academic.png",
        "dataset": ASSET_DIR / "dataset_overview_academic.png",
        "baseline_arch": ASSET_DIR / "baseline_architecture_academic.png",
        "finetune_arch": ASSET_DIR / "finetune_architecture_academic.png",
        "generation_flow": ASSET_DIR / "generation_pipeline_academic.png",
        "environment": ASSET_DIR / "training_environment_academic.png",
    }
    make_research_flow(values["flow"])
    make_baseline_figure(values["baseline"])
    make_training_figure(values["training"])
    make_generation_figure(values["generation"])
    make_dataset_figure(values["dataset"])
    make_baseline_architecture(values["baseline_arch"])
    make_finetune_architecture(values["finetune_arch"])
    make_generation_pipeline(values["generation_flow"])
    make_environment_figure(values["environment"])
    return values


def add_reference_list(document: Document) -> None:
    refs = [
        "[1] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems. 2020, 33: 9459-9474.",
        "[2] KARPUKHIN V, OĞUZ B, MIN S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]//Proceedings of EMNLP. 2020: 6769-6781.",
        "[3] ZHAI X, MUSTAFA B, KOLESNIKOV A, et al. Sigmoid Loss for Language Image Pre-Training[C]//Proceedings of ICCV. 2023: 11975-11986.",
        "[4] FAYSSE M, SIBILLE H, WU T, et al. ColPali: Efficient Document Retrieval with Vision Language Models[EB/OL]. arXiv:2407.01449, 2024.",
        "[5] SANTHANAM K, KHATTAB O, SAAD-FALCON J, et al. ColBERTv2: Effective and Efficient Retrieval via Lightweight Late Interaction[C]//Proceedings of NAACL. 2022: 3715-3734.",
        "[6] XIAO S, LIU Z, ZHANG P, et al. C-Pack: Packed Resources for General Chinese Embeddings[EB/OL]. arXiv:2309.07597, 2023.",
        "[7] MASRY A, LONG D X, TAN J Q, et al. ChartQA: A Benchmark for Question Answering about Charts with Visual and Logical Reasoning[C]//Findings of ACL. 2022: 2263-2279.",
        "[8] DU Y, LI C, GUO R, et al. PP-OCRv2: Bag of Tricks for Ultra Lightweight OCR System[EB/OL]. arXiv:2109.03144, 2021.",
        "[9] VAN DEN OORD A, LI Y, VINYALS O. Representation Learning with Contrastive Predictive Coding[EB/OL]. arXiv:1807.03748, 2018.",
        "[10] DOSOVITSKIY A, BEYER L, KOLESNIKOV A, et al. An Image is Worth 16×16 Words: Transformers for Image Recognition at Scale[C]//Proceedings of ICLR. 2021.",
        "[11] MARAFIOTI A, ZOHAR O, FARRÉ M, et al. SmolVLM: Redefining Small and Efficient Multimodal Models[EB/OL]. arXiv:2504.05299, 2025.",
        "[12] VIDORE. colSmol-500M Model Card[EB/OL]. https://huggingface.co/vidore/colSmol-500M, 2026-08-08.",
        "[13] HUGGING FACE. SmolVLM-500M-Instruct Model Card[EB/OL]. https://huggingface.co/HuggingFaceTB/SmolVLM-500M-Instruct, 2026-08-08.",
    ]
    for ref in refs:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(ref), size=10.5)


def build() -> None:
    figs = figures()
    document = Document()
    configure_document(document)
    configure_report_styles(document)

    # 封面：保持简洁，不放作者、学校、日期和装饰性说明。
    for _ in range(6):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    set_run_font(p.add_run(TITLE), east_asia="黑体", size=22, bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("项 目 报 告"), east_asia="楷体", size=18, bold=True)
    page_break(document)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("摘  要"), east_asia="黑体", size=18, bold=True)
    paragraph(document, (
        "针对复杂图表页面在文本化过程中容易丢失版式与视觉关系的问题，本项目围绕页面级视觉检索与检索增强问答开展模型设计和对照实验。"
        "项目从 ChartQA 中整理 2958 张去重图表页面和 6000 条真实问答，按 5000/500/500 划分训练集、验证集和测试集，并建立页面、Query、答案和证据页统一数据模式。"
        "检索部分设置三条基线：以 SigLIP 零样本单向量检索作为同架构基础对照，以 PP-OCR+BGE 代表传统 OCR 文本链路，以 ColSmol 代表多向量视觉晚交互方法；在此基础上，"
        "使用多正例 InfoNCE 对 SigLIP 文本塔和视觉塔末端层进行部分微调，并对 0、1、2 个困难负例进行消融。测试结果显示，部分微调把 SigLIP 的 Recall@10 从 0.2520 提升至 0.3240，"
        "MRR@10 从 0.1903 提升至 0.2320；ColSmol 获得最高 Recall@10=0.4460 和 MRR@10=0.3669；OCR+BGE 的 Recall@1=0.1980，说明文本链路在清晰图表标签上仍具有竞争力。"
        "困难负例没有进一步提高 MRR，反映出静态挖掘中假负例与训练构成变化的影响。生成部分使用 SmolVLM-500M-Instruct 比较 Oracle、Top-1、Top-3 逐页和 Top-3 拼接。"
        "ColSmol Top-3 逐页取得最高宽松准确率 0.1860，但 ColSmol Top-1 以 0.1840 的相近效果和 455.1 ms 的平均时延形成更好的成本折中。三个真实案例进一步表明，端到端错误分别可能来自检索缺失、"
        "生成模型算术错误和检索分数加权融合失效。项目最终完成了数据接入、三条检索基线、SigLIP 微调与消融、多候选页视觉问答以及可追溯评测的完整实验闭环。"
    ), font="楷体")
    paragraph(document, "关键词：视觉语言模型；页面检索；对比学习；SigLIP；ColSmol；检索增强生成", first_line=False, bold_prefix="关键词：")
    page_break(document)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("目  录"), east_asia="黑体", size=18, bold=True)
    add_toc(document)
    page_break(document)

    # 1
    add_heading(document, "1 研究背景与项目工作", 1)
    add_heading(document, "1.1 研究问题", 2)
    paragraph(document, "传统文档 RAG 通常先将 PDF 或扫描页转成文本，再按固定长度切片并建立文本向量索引。这种流程适合连续段落，却难以完整表达图表横纵轴、颜色、图例和数据点之间的关系。页面级视觉 RAG 保留完整页面图像，先通过视觉语言模型检索证据页，再由多模态生成模型读取证据并回答问题。它仍遵循“检索—增强—生成”的 RAG 结构[1]，但检索对象由文本 chunk 变为页面图像。")
    paragraph(document, "本项目关注的不是单纯调用一个大模型，而是研究四个相互关联的问题：不同页面表示方式会造成多大的检索差异；监督微调能否改善通用图文模型在图表页面上的排序；困难负例是否有效；检索召回提升后，Top-K 页面能否真正提高最终答案准确率。")

    add_heading(document, "1.2 本项目实际完成的工作", 2)
    paragraph(document, "数据接入方面，项目编写了 ChartQA 流式适配程序，将原始图像统一为 PNG，并依据图像内容生成稳定标识、完成去重和跨划分检查。整理后的 2958 张页面与 6000 条问答均保留 Query、标准答案和证据页之间的对应关系，最终形成 manifest.json、pages.jsonl 和 samples.jsonl 三类可复核数据文件。")
    paragraph(document, "检索实验方面，项目分别实现了 SigLIP 零样本、PP-OCR+BGE 和 ColSmol 三条完整链路，而不是只调用三个模型得到单个分数。每条链路都包含页面预处理、离线索引、Query 编码、Top-K 排序和统一指标计算，并保存逐条 Top-10 结果、运行时间与资源统计。")
    paragraph(document, "模型训练方面，项目在 SigLIP 基础上实现双塔末端层部分解冻、多正例 InfoNCE、混合精度、梯度累积与早停，得到 138 MB 检查点和四轮训练日志；随后从零样本排序中挖掘 1 个、2 个困难负例，完成同任务消融。生成阶段又实现 Oracle、Top-1、Top-3 逐页、分数加权融合和页面拼接，共保留 3000 条方法级预测及其时延和错误类型。")
    paragraph(document, "工程复现方面，训练配置、断点续跑、模型缓存、服务器迁移和结果归档均被整理为脚本。Stage 3 与 Stage 4 的指标、逐条预测、日志和检查点最终汇总进实验压缩包，使报告中的数字和案例可以回到原始文件核验。")
    analysis_paragraph(document, "以上工作中，SigLIP、BGE、ColSmol 和 SmolVLM 使用开源预训练权重；本项目的主要工作是数据规范、三条对照链路、SigLIP 任务微调、消融设置、Top-K 生成策略和统一评测。报告不会把预训练模型本身表述为本项目重新训练的成果。")

    add_heading(document, "1.3 研究主线", 2)
    add_figure(document, figs["flow"], "图 1-1 本项目以模型对照和实验验证为核心的研究主线", width=6.25)

    # 2
    add_heading(document, "2 数据集、任务定义与评价方法", 1)
    add_heading(document, "2.1 ChartQA 数据构建", 2)
    paragraph(document, "ChartQA 是面向图表视觉理解和逻辑推理的公开问答基准[7]。项目通过 HuggingFaceM4/ChartQA 的 streaming 接口读取原始 train、val 和 test，将 val 映射为 dev。每个图表经 RGB 转换后保存为 PNG，以 SHA-256 生成稳定 page_id；若同一页面出现在不同划分，仅保留第一次所属划分，避免页面图像泄漏。最终保留 2958 张唯一页面和 6000 条问答，其中训练 5000、验证 500、测试 500。")
    add_figure(document, figs["dataset"], "图 2-1 ChartQA 页面与问答样本的整理结果", width=6.05)
    add_table(document, ["对象", "字段与规模", "用途"], [
        ["页面", "page_id、image_path、doc_type、SHA-256；共 2958 页", "构成检索库和证据对象"],
        ["问答", "query_id、query、answers、evidence_page_ids；共 6000 条", "训练和端到端评估"],
        ["划分", "train/dev/test=5000/500/500", "训练、检查点选择和最终测试"],
        ["文档类型", "chart", "当前结论只针对图表页面"],
    ], widths=[3.0, 7.3, 5.6])

    add_heading(document, "2.2 任务定义", 2)
    paragraph(document, "给定 Query q 和页面库 P，检索器输出相关性分数 s(q,p)，并取前 K 个页面 P_K。生成器 g 读取问题和页面图像形成页面级候选答案；Top-3 逐页方案再通过融合函数 F 根据检索分数合并答案。")
    equation(document, "P_K = TopK₍p∈P₎ s(q,p)，   â = F({g(q,p) | p∈P_K}, {s(q,p)})")
    paragraph(document, "本实验中的“多页”指从全库检索出的多个候选图表页面，并不等同于同一个真实 PDF 内的跨页推理。该边界决定了报告只讨论候选页面组织和证据召回，不宣称已经解决企业长文档中的跨页关联问题。")

    add_heading(document, "2.3 评价指标", 2)
    add_table(document, ["指标", "计算含义", "在本项目中的判断作用"], [
        ["Recall@1/3/10", "证据页是否进入前 1/3/10 名", "判断候选覆盖范围"],
        ["MRR@10", "首个证据页倒数排名的均值", "同时考虑命中和排序位置"],
        ["EM", "规范化预测与标准答案完全一致", "严格问答指标"],
        ["Relaxed Accuracy", "精确匹配或数值误差不超过 5%", "ChartQA 主分析指标"],
        ["Accuracy given hit", "已召回证据页样本中的宽松准确率", "隔离观察生成能力"],
        ["Mean generation ms", "单个方法平均生成时延", "比较效果与计算成本"],
    ], widths=[3.4, 6.3, 6.0])

    add_heading(document, "2.4 训练平台与软件环境", 2)
    paragraph(document, "正式实验运行在 AutoDL 云端容器中。基础镜像为 Ubuntu 22.04、Python 3.12、PyTorch 2.8.0 和 CUDA 12.8，NVIDIA 驱动报告的最高 CUDA 兼容版本为 13.0；这里的 13.0 是驱动能力，实际 PyTorch 运行时仍按镜像中的 CUDA 12.8 构建。核心依赖包括 Transformers 5.14.1、Datasets 5.0.1、Sentence-Transformers 5.6.1、PaddleOCR 3.7.0 和 colpali-engine 0.3.16。")
    paragraph(document, "SigLIP 零样本评估和部分微调首先在单张 RTX 5090 32 GB 上完成，实例配置同时提供 25 核 CPU 和 92 GB 内存。由于后续时段 5090 无空闲卡，项目将数据盘中的数据集、Hugging Face 缓存、索引和检查点迁移到单张 RTX 4090 24 GB 实例，继续完成 PP-OCR+BGE、ColSmol 和 Stage 4 视觉问答。迁移只改变硬件实例，不修改代码和模型参数。")
    paragraph(document, "数据与模型缓存统一存放于 /root/autodl-tmp，避免占用较小的系统盘。SigLIP 训练使用 bfloat16 自动混合精度、梯度累积和随机种子 42；正式端到端方法的时延均在同一 4090 环境下计算，因此第 6.4 节各方法可以直接比较。OCR 抽取属于独立预处理，耗时不计入生成阶段的平均生成时延。")
    add_figure(document, figs["environment"], "图 2-2 实际训练环境与云端实例迁移过程", width=6.1)

    # 3: three baselines
    add_heading(document, "3 三条检索基线的模型设计与作用", 1)
    paragraph(document, "三条基线不是三个随意选择的模型，而是分别代表三种页面信息处理假设：全页压缩为一个视觉向量、先 OCR 再做文本向量、保留页面局部 token 并进行晚交互。三者共同回答“视觉信息是否必要、单向量是否足够、OCR 文本链路是否仍有价值”。微调 SigLIP 则在第一条基线的同一架构上进行任务适配，因此它属于本项目训练方法而不是第四条独立基线。")
    add_figure(document, figs["baseline_arch"], "图 3-1 三条检索基线的数据流与页面表示差异", width=6.2)

    add_heading(document, "3.1 基线一：SigLIP 零样本单向量检索", 2)
    paragraph(document, "SigLIP 由文本编码器和视觉编码器构成，预训练时对图文对使用 sigmoid 损失进行对齐[3]。本项目加载 google/siglip-base-patch16-224，输入 Query 和 224×224 页面图像，分别调用文本与图像特征接口，输出 L2 归一化的全局向量。检索得分为向量内积，即归一化后的余弦相似度。")
    equation(document, "q = f_text(Query)/‖f_text(Query)‖₂，vᵢ = f_image(Pageᵢ)/‖f_image(Pageᵢ)‖₂，sᵢ=qᵀvᵢ")
    paragraph(document, "这条基线的作用有两点。第一，它完全不依赖 OCR，可以验证通用图文预训练模型能否直接用于页面检索；第二，它与后续部分微调使用相同基础模型、相同输入和相同检索规则，因此二者差值可以较为直接地归因于任务监督。其缺点是整页只保留一个向量，图表内多个标签、数值和颜色在池化后可能相互干扰。")

    add_heading(document, "3.2 基线二：PP-OCR+BGE 文本检索", 2)
    paragraph(document, "OCR 基线先使用 PP-OCRv6-medium 检测并识别页面文字，过滤置信度低于 0.5 的文本行，再把同页文本拼接为一个文档字符串。Query 与页面 OCR 文本由 BAAI/bge-small-en-v1.5 编码，向量归一化后进行余弦检索。BGE 属于面向检索的通用文本嵌入模型，其相关资源覆盖多任务文本表示训练[6]。")
    paragraph(document, "选择该基线是为了回答“直接视觉检索是否比传统 OCR-RAG 更合适”。它的优势是图表标题和标签一旦识别正确，文本模型可以精确匹配实体词；劣势是柱形长度、颜色、坐标位置等视觉关系无法完整进入文本向量，而且 OCR 增加了独立的离线处理阶段。实验中 2958 页均产生非空 OCR 文本，平均单页处理时间为 7571.2 ms，因此该基线不是故意削弱的空文本对照。")

    add_heading(document, "3.3 基线三：ColSmol 多向量晚交互检索", 2)
    paragraph(document, "ColSmol 使用 vidore/colSmol-500M，由 ColIdefics3 视觉语言骨干与检索适配器生成多向量表示。页面不被压缩为单个向量，而是保留多个视觉 token；Query 也保留 token 级表示，最后通过 score_multi_vector 完成 MaxSim 风格晚交互。ColPali 类工作表明，直接对文档页面图像构造多向量并做晚交互可以利用文本、布局和图像线索[4]；其机制与 ColBERT 的 token 级晚交互思想相近[5]。")
    paragraph(document, "这条基线代表更强的视觉检索方法。对于 Query 中的国家名、年份、颜色或图表对象，不同 token 可以分别匹配页面的不同局部区域，因此不会像单向量那样过早合并信息。代价是每页保存多个向量，页面编码和相似度计算更复杂。本实验中页面编码耗时 1168.9 s，6000 条 Query 的编码与全库评分耗时 27.1 s，峰值 GPU 显存约 5605 MB。")

    add_heading(document, "3.4 三条基线的可比关系", 2)
    paragraph(document, "SigLIP 和 ColSmol 都直接读取整页图像，但前者将页面压缩成一个全局向量，后者保留视觉 token 级多向量；OCR+BGE 则先丢弃视觉布局，把识别文字压缩成一个文本向量。SigLIP 与 OCR+BGE 最终都用余弦相似度排序，ColSmol 使用晚交互评分。三条基线都保持预训练权重不变，因此分别承担同架构基础对照、传统文本链路对照和强视觉检索对照的作用。")
    analysis_paragraph(document, "三条基线的输入和表示粒度并不完全相同，因此结果用于比较完整方法链路，不用于声称同参数量下某种架构必然优于另一种架构。只有“SigLIP 零样本—SigLIP 部分微调”属于较严格的同架构对照。")

    # 4: finetune
    add_heading(document, "4 SigLIP 检索器训练与困难负例消融", 1)
    add_heading(document, "4.1 部分微调策略", 2)
    paragraph(document, "训练脚本首先冻结 SigLIP 全部参数，再解冻文本编码器最后两层、视觉编码器最后两层以及两塔的归一化与投影头。总参数量为 203,155,970，可训练参数为 36,032,256，占 17.74%。选择部分微调而不是全量训练，是为了在单张 32 GB GPU 上保留预训练图文对齐能力，同时让高层语义适配 ChartQA 的 Query—页面关系。")
    add_figure(document, figs["finetune_arch"], "图 4-1 SigLIP 双塔部分微调结构与 InfoNCE 训练目标", width=6.05)
    add_table(document, ["训练项", "配置"], [
        ["基础模型", "google/siglip-base-patch16-224"], ["Epoch / Batch", "4 / 16"],
        ["梯度累积 / Eval Batch", "2 / 32"], ["学习率 / 权重衰减", "2×10⁻⁶ / 0.01"],
        ["温度 / Warmup", "0.07 / 0.05"], ["解冻范围", "文本末 2 层 + 视觉末 2 层 + 两塔头部"],
        ["优化与精度", "AdamW、bfloat16、梯度裁剪 1.0"], ["模型选择", "验证集 MRR@10，patience=2，seed=42"],
    ], widths=[6.2, 9.0])

    add_heading(document, "4.2 多正例 InfoNCE", 2)
    paragraph(document, "一个批次内可能有多个 Query 指向同一图表页面。实现没有强制一一配对，而是根据 page_id 构造正例掩码，使同页的所有 Query—Page 组合都作为正例，其余页面作为批内负例。对第 i 个 Query，损失为：")
    equation(document, "Lᵢ=−log[Σⱼ∈P(i)exp(sᵢⱼ/τ) / Σₖexp(sᵢₖ/τ)]，τ=0.07")
    paragraph(document, "训练采用自动混合精度和梯度累积，每个 Epoch 后重新编码验证集全部页面和 Query，以 MRR@10 选择最佳状态。保存的检查点只包含解冻层参数和基础模型信息，实际使用时在原始 SigLIP 权重上恢复。")

    add_heading(document, "4.3 困难负样本构造", 2)
    paragraph(document, "困难负例来自零样本 SigLIP 的 Top-10 检索结果。对每条训练 Query，剔除 evidence_page_ids 后，依次选取排名最高的 1 个或 2 个页面加入当前 batch。这些页面比随机批内负例更相似，理论上能迫使模型学习更细的区分边界。项目分别训练 hardneg_1 和 hardneg_2，与不额外加入困难负例的部分微调形成消融。")
    analysis_paragraph(document, "该消融不是完全无混杂变量。普通微调、1 个困难负例和 2 个困难负例的单步 batch size 分别为 16、12、8，虽然梯度累积后有效批量接近，但每个 Query 实际参与对比的页面构成不同。因此结果应解释为“当前完整训练配置”的对比，而不是困难负例数量的纯因果效应。")

    # 5 generation
    add_heading(document, "5 检索结果驱动的视觉问答", 1)
    add_heading(document, "5.1 生成模型与提示词", 2)
    paragraph(document, "生成器使用 HuggingFaceTB/SmolVLM-500M-Instruct。该模型属于轻量视觉语言模型，可处理图像与文本并生成答案[11]。项目使用统一提示词要求模型仅依据图表图像回答，并只返回短答案；生成采用 do_sample=False，最大生成 16 token。SmolVLM 在本项目中没有重新训练，其作用是固定生成器条件下比较不同证据页组织方式。")

    add_heading(document, "5.2 四类问答设置", 2)
    paragraph(document, "Oracle 直接把标准 evidence page 交给生成器，用来测量正确证据条件下的生成上限。Top-1 只读取检索首位页面，是最短的检索—生成闭环。Top-3 逐页对前三页分别生成答案，再按检索得分融合三个候选；Top-3 拼接则将三页纵向缩放为一张图，只调用一次生成器，用于观察速度收益和图像压缩损失。")
    add_figure(document, figs["generation_flow"], "图 5-1 Top-1、Top-3 逐页与 Top-3 拼接的处理流程", width=6.15)

    add_heading(document, "5.3 检索分数加权融合", 2)
    paragraph(document, "逐页方案先对三个检索分数执行 softmax，再按规范化后的答案字符串聚合权重。相同答案的页面权重相加，权重最大的答案成为最终结果。")
    equation(document, "wᵢ=exp(sᵢ/T)/Σⱼexp(sⱼ/T)，Score(a)=Σᵢ 𝟙[norm(aᵢ)=norm(a)]·wᵢ")
    paragraph(document, "该方法把检索相关性当作候选答案置信度，计算简单且可追溯，但检索得分只表示 Query 与页面相关，不直接表示页面答案是否正确。真实案例将说明，当证据页排在第二位且分差很小时，Top-3 虽然召回正确页面，最终仍可能被第一候选页的错误答案主导。")

    # 6 Results
    add_heading(document, "6 实验结果与真实案例分析", 1)
    add_heading(document, "6.1 训练曲线与最佳检查点", 2)
    add_figure(document, figs["training"], "图 6-1 SigLIP 部分微调训练曲线", width=5.75)
    add_table(document, ["Epoch", "Train Loss", "Dev\nRecall@1", "Dev\nRecall@3", "Dev\nRecall@10", "Dev\nMRR@10"], [
        ["1", "2.0723", "0.2100", "0.2780", "0.3440", "0.2511"],
        ["2", "1.6189", "0.2100", "0.2840", "0.3500", "0.2551"],
        ["3", "1.4352", "0.2140", "0.2860", "0.3580", "0.2591"],
        ["4", "1.3690", "0.2140", "0.2840", "0.3580", "0.2588"],
    ], widths=[2.0, 3.0, 2.5, 2.5, 2.7, 3.1], font_size=10)
    analysis_paragraph(document, "训练损失由 2.0723 降至 1.3690，下降约 33.9%；验证集 MRR@10 在第 3 轮达到 0.2591，第 4 轮略降至 0.2588，因此检查点选择第 3 轮。损失继续下降而排序指标不再上升，说明继续拟合训练对比目标并不等价于提升验证集检索质量。")

    add_heading(document, "6.2 三条基线与微调方法的检索结果", 2)
    add_table(document, ["排名", "方法", "Recall@1", "Recall@3", "Recall@10", "MRR@10"], [
        ["1", "ColSmol 基线", "0.3320", "0.3900", "0.4460", "0.3669"],
        ["2", "SigLIP 部分微调", "0.1900", "0.2600", "0.3240", "0.2320"],
        ["3", "PP-OCR+BGE 基线", "0.1980", "0.2540", "0.2800", "0.2274"],
        ["4", "SigLIP 零样本基线", "0.1600", "0.2160", "0.2520", "0.1903"],
    ], widths=[1.6, 5.7, 2.2, 2.2, 2.5, 2.7], font_size=10)
    add_figure(document, figs["baseline"], "图 6-2 三条检索基线与 SigLIP 微调方法的测试结果", width=6.05)
    paragraph(document, "同架构比较中，部分微调相对 SigLIP 零样本使 R@1、R@3、R@10 和 MRR@10 分别提高 0.0300、0.0440、0.0720 和 0.0416；相对增幅为 18.75%、20.37%、28.57% 和 21.88%。这证明使用 Query—证据页监督信号更新两塔末端层是有效的。")
    paragraph(document, "OCR+BGE 的 R@1=0.1980，反而略高于微调 SigLIP 的 0.1900，说明图表标题和标签能被正确 OCR 时，文本实体匹配仍然有优势。但 OCR+BGE 的 R@10=0.2800、MRR@10=0.2274 均低于部分微调，表明它对部分问题非常准确，却没有形成更稳定的候选覆盖。")
    paragraph(document, "ColSmol 在四项指标上均排名第一，R@10 比微调 SigLIP 高 0.1220，MRR@10 高 0.1349。这个差距说明图表页面包含大量局部匹配信号，多向量晚交互比单个全局向量更适合当前任务。与此同时，ColSmol 是更强的预训练视觉检索模型，结果不应表述为本项目微调方法超过所有基线；实际结论是：部分微调改善了轻量单向量模型，而强晚交互基线仍保持明显领先。")

    add_heading(document, "6.3 困难负例消融", 2)
    add_table(document, ["方法", "Recall@1", "Recall@3", "Recall@10", "MRR@10"], [
        ["SigLIP 零样本", "0.1600", "0.2160", "0.2520", "0.1903"],
        ["部分微调（0 个困难负例）", "0.1900", "0.2600", "0.3240", "0.2320"],
        ["部分微调（1 个困难负例）", "0.1840", "0.2600", "0.3240", "0.2313"],
        ["部分微调（2 个困难负例）", "0.1660", "0.2640", "0.3240", "0.2215"],
    ], widths=[6.6, 2.4, 2.4, 2.7, 2.7], font_size=10)
    analysis_paragraph(document, "增加困难负例没有提高 MRR。1 个负例使 MRR 轻微下降 0.0007；2 个负例下降 0.0105，同时 R@1 从 0.1900 降到 0.1660。R@10 始终为 0.3240，说明负例主要改变证据页在前十名内部的位置，而没有扩大命中集合。可能原因包括静态零样本排名产生假负例、相似图表之间监督不充分，以及训练 batch 构成变化。")

    add_heading(document, "6.4 端到端问答结果", 2)
    add_table(document, ["方法", "证据\n召回率", "Exact\nMatch", "Relaxed\nAccuracy", "Accuracy\ngiven hit", "平均生成\n时延/ms"], [
        ["Oracle", "1.0000", "0.3180", "0.3460", "0.3460", "466.3"],
        ["SigLIP Top-1", "0.1900", "0.1280", "0.1620", "0.4632", "465.0"],
        ["SigLIP Top-3 逐页", "0.2600", "0.1280", "0.1640", "0.3462", "1387.0"],
        ["ColSmol Top-1", "0.3320", "0.1540", "0.1840", "0.3675", "455.1"],
        ["ColSmol Top-3 逐页", "0.3900", "0.1580", "0.1860", "0.3333", "1369.5"],
        ["ColSmol Top-3 拼接", "0.3900", "0.0980", "0.1260", "0.2000", "350.3"],
    ], widths=[4.4, 2.7, 1.8, 2.8, 3.0, 2.5], font_size=9.5)
    add_figure(document, figs["generation"], "图 6-3 端到端问答准确率与平均生成时延", width=6.1)
    analysis_paragraph(document, "ColSmol Top-3 逐页的宽松准确率最高，但只比 Top-1 高 0.0020，时延却从 455.1 ms 增至 1369.5 ms。Top-3 拼接虽然降到 350.3 ms，但准确率比逐页方案低 0.0600。Oracle 仍只有 0.3460，说明当前端到端结果同时受检索器和 SmolVLM-500M 的图表读取、数值计算能力限制。")

    add_heading(document, "6.5 真实案例一：证据页正确，生成模型计算错误", 2)
    case_info(document, "What is the difference in value between Lamb and Corn?", "0.57", "chartqa_16ed5f0796ef0b10（ChartQA test row 1）")
    add_figure(document, CASE_DIR / "case_generation_error.png", "图 6-4 案例一的真实 ChartQA 证据页面", width=5.7)
    add_table(document, ["方法", "证据是否命中", "模型输出", "判断"], [
        ["Oracle", "是", "1.07", "错误"],
        ["SigLIP Top-1", "否", "0.15", "检索错误"],
        ["ColSmol Top-1", "是，证据页排第 1", "1.07", "生成错误"],
    ], widths=[4.1, 5.0, 3.0, 3.0], font_size=10)
    paragraph(document, "图中 Lamb=103.70，Corn=103.13，差值应为 0.57。ColSmol 已把标准证据页排到第一位，但 SmolVLM 输出 1.07；Oracle 也得到相同错误，因此该样本不是检索失败，而是生成模型的数值读取或减法错误。它直接说明 Oracle 0.3460 所反映的生成上限问题。")

    add_heading(document, "6.6 真实案例二：ColSmol 的局部视觉匹配优势", 2)
    case_info(document, "What is the value of Slovenia in the graph?", "1", "chartqa_c41db587bcf088b9（ChartQA test row 10）")
    add_figure(document, CASE_DIR / "case_colsmol_win.png", "图 6-5 案例二的真实 ChartQA 证据页面", width=5.65)
    add_table(document, ["方法", "Top-1 页面", "输出", "结果"], [
        ["SigLIP 部分微调", "chartqa_ef30b6391fd594aa（非证据页）", "0.22", "错误"],
        ["ColSmol Top-1", "chartqa_c41db587bcf088b9（证据页）", "1.00", "宽松正确"],
        ["ColSmol Top-3 逐页", "证据页排第 1，另两页输出 12.6、60", "1.00", "宽松正确"],
        ["ColSmol Top-3 拼接", "相同三页", "1.", "当前指标判错"],
    ], widths=[4.1, 6.3, 2.8, 2.6], font_size=9.5)
    paragraph(document, "该页面同时包含国家名、水平柱、数值标签和解释文字。SigLIP 的全局向量把一个非证据图表排在首位，而 ColSmol 通过局部 token 晚交互命中包含“Slovenia”和“1 GPI”的页面，体现了第三条基线的实际优势。另一个值得注意的问题是拼接输出“1.”被当前规范化函数判错：字符规范化保留小数点，而数值解析又不接受末尾没有数字的小数点。这个案例暴露了评测实现对格式的敏感性。")

    add_heading(document, "6.7 真实案例三：Top-3 已召回证据，但加权融合仍失败", 2)
    case_info(document, "What's the percentage of U.S adults who refused?", "2", "chartqa_ad4f88cf12770f47（ChartQA test row 14）")
    add_figure(document, CASE_DIR / "case_fusion_error.png", "图 6-6 案例三的真实 ChartQA 证据页面", width=4.5)
    add_table(document, ["方法/候选", "页面与得分", "候选答案", "最终结果"], [
        ["SigLIP Top-1", "证据页，score=0.1746", "2%", "正确"],
        ["ColSmol Top-1", "非证据页，score=16.8415", "28%", "错误"],
        ["ColSmol Top-3 第 2 页", "证据页，score=16.7906", "2%", "已召回"],
        ["ColSmol Top-3 融合", "28% 权重 0.3997；2% 权重 0.3799；1% 权重 0.2204", "28%", "错误"],
    ], widths=[4.0, 6.7, 2.5, 2.5], font_size=9.4)
    paragraph(document, "ColSmol Top-3 已经把标准证据页召回到第二位，但第一候选页的检索得分略高，softmax 后 28% 的权重为 0.3997，高于正确答案 2% 的 0.3799，最终仍选择错误答案。该案例解释了为什么 ColSmol 从 Top-1 到 Top-3 的检索召回提高 0.0580，而宽松准确率只提高 0.0020：检索命中并不保证简单分数投票能识别正确候选答案。")

    # 7 implementation/work
    add_heading(document, "7 代码实现与实验产出", 1)
    add_heading(document, "7.1 核心代码及其在数据流中的位置", 2)
    paragraph(document, "代码并不是彼此独立的功能清单，而是沿着“数据整理—页面编码—检索训练—视觉生成—结果汇总”顺序连接。下面按照实际运行先后说明核心文件的输入、处理和输出。")
    code_module(document, "dataset_adapters.py 与 dataset_schema.py", "前者读取 ChartQA 原始样本并生成稳定 page_id，后者定义 DocumentPage 与 DocumentQASample 的统一结构。二者把零散图像和问答转成 pages.jsonl、samples.jsonl 与页面目录，是所有训练、索引和评测脚本共同的数据入口。")
    code_module(document, "siglip_encoder.py 与 siglip_index.py", "编码器分别接收页面图像和 Query，输出 L2 归一化的 SigLIP 全局向量；索引模块把页面向量及其 page_id 持久化，并利用矩阵内积返回 Top-K。零样本基线和微调检查点评估都复用这条检索接口。")
    code_module(document, "ocr_extractor.py 与 evaluate_ocr_retrieval.py", "OCR 模块调用 PP-OCR 对每页执行文字检测和识别，过滤低置信度文本行；评估脚本随后用 BGE 编码 OCR 文本和 Query，建立文本向量索引并计算 Recall 与 MRR。其输出既是第二条基线结果，也记录逐页 OCR 耗时和空文本统计。")
    code_module(document, "colsmol_encoder.py 与 evaluate_colsmol_retrieval.py", "页面和 Query 在这里被编码为多组 token 向量，不再压缩为单个向量。评估脚本分批加载页面表示，执行晚交互评分并保存 Top-10 排序，同时记录页面编码时间、Query 评分时间和峰值显存。")
    code_module(document, "train_siglip_partial_finetune.py", "该文件承担项目中真正的检索器训练。它先冻结基础权重，再定位并解冻文本塔和视觉塔末端两层，构造多正例掩码计算 InfoNCE，并负责 bfloat16 混合精度、梯度累积、早停、困难负例读取以及最佳检查点保存。")
    code_module(document, "evaluate_siglip_partial_checkpoint.py", "训练完成后，该脚本恢复最佳解冻层参数，重新编码 2958 张页面与测试 Query，而不是直接沿用训练时缓存；随后按统一格式输出 retrieval_results.json 和 metrics.json，保证微调结果能够与三条基线直接比较。")
    code_module(document, "smolvlm_generator.py", "该模块把 Query 和页面图像组织为 SmolVLM 对话输入，统一图像预处理、短答案提示词、最大生成长度与确定性解码。它只封装开源生成模型推理，没有对 SmolVLM 再训练。")
    code_module(document, "evaluate_stage4_generation.py", "这是检索与生成闭环的主程序。它读取 SigLIP 或 ColSmol 的 Top-K 结果，执行 Oracle、Top-1、Top-3 逐页和拼接实验；逐页答案可缓存续跑，最终由分数加权模块融合，并输出 EM、宽松准确率、命中条件准确率、时延和逐条错误记录。")
    code_module(document, "summarize_*.py 与 run_stage*.sh", "前者从各实验目录读取 JSON 并生成统一 Markdown、CSV 对比结果，后者把依赖安装、索引、评测、训练和归档命令连接成可重复流水线。服务器中断或迁移后，可以依据已有文件跳过已完成步骤，而不必重新运行全部实验。")

    add_heading(document, "7.2 最终实验文件", 2)
    paragraph(document, "阶段三归档包含 SigLIP 零样本、部分微调、两个困难负例变体、OCR+BGE、ColSmol 的 metrics.json、retrieval_results.json、训练日志和检查点；阶段四归档包含 run_config.json、page_answers.jsonl、collage_answers.jsonl、predictions.json、metrics.json 与对比表。逐条预测记录 Query、标准答案、证据页、候选页、模型输出、时延和错误类型，使本报告的统计表和真实案例都能够追溯到原始实验。")
    analysis_paragraph(document, "本项目已经完成模型研究所需的主要闭环：统一数据、基础对照、强基线、同架构微调、负例消融、端到端评估和案例诊断。报告中的核心贡献不是提出一个全新的预训练 VLM，而是在同一真实数据集上把不同视觉/文本检索范式和生成策略做成了可运行、可比较、可复核的工程实验。")

    # 8 conclusion
    add_heading(document, "8 结论", 1)
    paragraph(document, "本项目围绕页面级视觉检索与问答完成了以模型和实验为中心的研究。三条基线分别覆盖通用单向量视觉检索、传统 OCR 文本检索和多向量晚交互视觉检索，使不同页面表示假设能够在同一 ChartQA 测试集上比较。SigLIP 部分微调相对其零样本基线显著提高检索指标，证明领域 Query—证据页监督有效；困难负例消融没有带来额外收益，说明负例筛选质量和训练构成需要谨慎控制。")
    paragraph(document, "ColSmol 获得最好的检索结果，表明图表页面受益于局部视觉 token 级匹配。端到端实验同时说明检索指标不是最终答案的充分条件：Top-3 提升了证据召回，却只带来极小的准确率增益；图片拼接降低时延但损害阅读效果；Oracle 条件下的生成错误则暴露了轻量 VLM 的图表读取和计算限制。真实案例进一步把总体指标还原为三种具体机制：检索器选错页面、生成器读对页面但算错答案、证据已进入 Top-3 但融合规则选择错误候选。")
    paragraph(document, "因此，本项目得到的主要结论是：页面级视觉 RAG 的性能由页面表示、检索排序、候选组织和视觉生成共同决定；对当前 ChartQA 规模而言，ColSmol Top-1 在准确率和时延之间形成较好的实际折中，而 SigLIP 部分微调提供了一条成本更低、效果可通过监督数据稳定改善的单向量路线。")

    add_heading(document, "参考文献", 1)
    add_reference_list(document)

    # Remove running header text and retain page numbers only.
    add_page_numbers_only(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
