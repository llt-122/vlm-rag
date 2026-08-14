from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "docs" / "report_assets"
OUTPUT = ROOT / "docs" / "VLM-RAG项目报告_正式版.docx"


TITLE = "基于视觉语言模型的页面级文档检索增强问答系统设计与实验"


def set_run_font(run, east_asia: str = "宋体", western: str = "Times New Roman", size: float = 12,
                 bold: bool = False, color: str = "000000") -> None:
    run.font.name = western
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=9)


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在打开 Word 时自动更新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])
    set_run_font(run, size=12)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.extend([r_pr, text_node])
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_body(document: Document, text: str, *, first_line: bool = True, bold_prefix: str | None = None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = document.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    if first_line:
        pf.first_line_indent = Cm(0.85)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, east_asia="黑体", size=12, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=12)
    else:
        run = p.add_run(text)
        set_run_font(run, size=12)


def add_equation(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, east_asia="宋体", western="Cambria Math", size=11.5)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_heading(document: Document, text: str, level: int) -> None:
    p = document.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12}
    set_run_font(run, east_asia="黑体", size=sizes[level], bold=True)


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=10.5)


def add_figure(document: Document, path: Path, caption: str, width: float = 6.2) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(document, caption)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None,
              font_size: float = 10.5) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, east_asia="黑体", size=font_size, bold=True)
        if widths:
            cell.width = Cm(widths[idx])
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            set_run_font(run, size=font_size)
            if widths:
                cell.width = Cm(widths[idx])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.3)
    section.footer_distance = Cm(1.3)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for level, size in ((1, 16), (2, 14), (3, 12)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_headers_and_footers(document: Document) -> None:
    for section in document.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run("页面级文档 VLM-RAG 项目报告")
        set_run_font(hr, size=9, color="666666")
        add_page_field(section.footer.paragraphs[0])


def setup_plot_font() -> None:
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    if font_path.exists():
        font_manager.fontManager.addfont(str(font_path))
        plt.rcParams["font.family"] = font_manager.FontProperties(fname=str(font_path)).get_name()
    plt.rcParams["axes.unicode_minus"] = False


def make_architecture_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11.8, 5.4))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)
    ax.axis("off")

    ax.add_patch(FancyBboxPatch((0.15, 3.25), 11.7, 2.35, boxstyle="round,pad=0.03", fc="#EFF6FF", ec="#8DB3E2", lw=1.2))
    ax.add_patch(FancyBboxPatch((0.15, 0.35), 11.7, 2.35, boxstyle="round,pad=0.03", fc="#F4FBF4", ec="#93C47D", lw=1.2))
    ax.text(0.42, 5.25, "离线索引链路", fontsize=13, fontweight="bold", color="#24527A")
    ax.text(0.42, 2.35, "在线问答链路", fontsize=13, fontweight="bold", color="#38761D")

    def box(x, y, w, h, text, color):
        patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08", fc="white", ec=color, lw=1.4)
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=10.5)

    def arrow(x1, y1, x2, y2, color="#666666"):
        ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=13, lw=1.3, color=color))

    offline = [(0.6, "文档/图表\n页面图像"), (3.0, "统一数据模式\nPage + QA"), (5.4, "视觉编码器\nSigLIP/ColSmol"), (7.8, "页面向量或\n多向量表示"), (10.2, "持久化\n页面索引")]
    for x, label in offline:
        box(x, 3.85, 1.55, 0.9, label, "#5B9BD5")
    for x in (2.15, 4.55, 6.95, 9.35):
        arrow(x, 4.3, x + 0.75, 4.3, "#5B9BD5")

    online = [(0.6, "用户 Query"), (2.55, "文本编码"), (4.5, "相似度检索\nTop-K"), (6.45, "候选页面\n逐页/拼接"), (8.4, "SmolVLM\n视觉问答"), (10.35, "答案 +\n证据页面")]
    for x, label in online:
        box(x, 0.95, 1.35, 0.9, label, "#70AD47")
    for x in (1.95, 3.9, 5.85, 7.8, 9.75):
        arrow(x, 1.4, x + 0.5, 1.4, "#70AD47")

    arrow(10.95, 3.85, 5.2, 1.85, "#A55A20")
    ax.text(8.3, 2.75, "读取索引并返回候选页", fontsize=9.5, color="#A55A20", rotation=-9)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_training_figure(path: Path) -> None:
    epochs = [1, 2, 3, 4]
    losses = [2.072257, 1.618936, 1.435246, 1.368994]
    mrr = [0.251079, 0.255093, 0.259061, 0.258815]
    fig, ax1 = plt.subplots(figsize=(8.8, 4.7))
    ax2 = ax1.twinx()
    line1 = ax1.plot(epochs, losses, marker="o", lw=2.2, color="#4472C4", label="训练损失")
    line2 = ax2.plot(epochs, mrr, marker="s", lw=2.2, color="#ED7D31", label="验证集 MRR@10")
    ax1.set_xlabel("训练轮次（Epoch）")
    ax1.set_ylabel("平均训练损失", color="#4472C4")
    ax2.set_ylabel("验证集 MRR@10", color="#ED7D31")
    ax1.set_xticks(epochs)
    ax1.grid(axis="y", alpha=0.25)
    ax1.set_ylim(1.25, 2.2)
    ax2.set_ylim(0.245, 0.263)
    lines = line1 + line2
    ax1.legend(lines, [line.get_label() for line in lines], loc="center right")
    ax1.set_title("SigLIP 部分微调训练过程")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_retrieval_figure(path: Path) -> None:
    methods = ["SigLIP\n零样本", "OCR+BGE", "SigLIP\n部分微调", "ColSmol"]
    r1 = [0.160, 0.198, 0.190, 0.332]
    r3 = [0.216, 0.254, 0.260, 0.390]
    r10 = [0.252, 0.280, 0.324, 0.446]
    mrr = [0.1903, 0.2274, 0.2320, 0.3669]
    x = list(range(len(methods)))
    width = 0.19
    fig, ax = plt.subplots(figsize=(9.4, 5.0))
    for offset, values, label, color in ((-1.5, r1, "Recall@1", "#5B9BD5"), (-0.5, r3, "Recall@3", "#70AD47"), (0.5, r10, "Recall@10", "#FFC000"), (1.5, mrr, "MRR@10", "#ED7D31")):
        bars = ax.bar([v + offset * width for v in x], values, width, label=label, color=color)
        ax.bar_label(bars, fmt="%.3f", fontsize=8, padding=2)
    ax.set_xticks(x, methods)
    ax.set_ylim(0, 0.52)
    ax.set_ylabel("指标值")
    ax.set_title("测试集页面检索结果对比")
    ax.grid(axis="y", alpha=0.22)
    ax.legend(ncol=4, loc="upper left")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_generation_figure(path: Path) -> None:
    labels = ["Oracle", "SigLIP\nTop-1", "SigLIP\nTop-3逐页", "ColSmol\nTop-1", "ColSmol\nTop-3逐页", "ColSmol\nTop-3拼接"]
    acc = [0.346, 0.162, 0.164, 0.184, 0.186, 0.126]
    latency = [466.3, 465.0, 1387.0, 455.1, 1369.5, 350.3]
    x = list(range(len(labels)))
    fig, ax1 = plt.subplots(figsize=(10.4, 5.0))
    bars = ax1.bar(x, acc, color=["#A5A5A5", "#5B9BD5", "#4472C4", "#70AD47", "#548235", "#FFC000"], width=0.62, label="Relaxed Accuracy")
    ax1.bar_label(bars, fmt="%.3f", fontsize=8.5, padding=2)
    ax1.set_ylim(0, 0.40)
    ax1.set_ylabel("宽松准确率")
    ax1.set_xticks(x, labels)
    ax1.grid(axis="y", alpha=0.2)
    ax2 = ax1.twinx()
    line = ax2.plot(x, latency, color="#C00000", marker="o", lw=2.0, label="平均生成时延")
    ax2.set_ylabel("平均生成时延（ms）", color="#C00000")
    ax2.set_ylim(0, 1600)
    handles = [bars, line[0]]
    ax1.legend(handles, ["宽松准确率", "平均生成时延"], loc="upper right")
    ax1.set_title("不同检索—生成组合的准确率与时延")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_error_figure(path: Path) -> None:
    methods = ["SigLIP\nTop-1", "SigLIP\nTop-3逐页", "ColSmol\nTop-1", "ColSmol\nTop-3逐页", "ColSmol\nTop-3拼接"]
    correct = [81, 82, 92, 93, 63]
    retrieval_miss = [368, 333, 303, 277, 281]
    generation_error = [51, 85, 105, 130, 156]
    x = list(range(len(methods)))
    fig, ax = plt.subplots(figsize=(9.6, 4.8))
    ax.bar(x, correct, label="回答正确", color="#70AD47")
    ax.bar(x, retrieval_miss, bottom=correct, label="检索未命中且回答错误", color="#ED7D31")
    bottoms = [a + b for a, b in zip(correct, retrieval_miss)]
    ax.bar(x, generation_error, bottom=bottoms, label="已命中但生成错误", color="#A5A5A5")
    ax.set_xticks(x, methods)
    ax.set_ylim(0, 550)
    ax.set_ylabel("测试问题数（共 500 条）")
    ax.set_title("端到端错误来源分解")
    ax.legend(ncol=3, loc="upper center")
    ax.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_figures() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    setup_plot_font()
    paths = {
        "architecture": ASSET_DIR / "system_architecture.png",
        "training": ASSET_DIR / "training_curve.png",
        "retrieval": ASSET_DIR / "retrieval_comparison.png",
        "generation": ASSET_DIR / "generation_tradeoff.png",
        "errors": ASSET_DIR / "error_decomposition.png",
    }
    make_architecture_figure(paths["architecture"])
    make_training_figure(paths["training"])
    make_retrieval_figure(paths["retrieval"])
    make_generation_figure(paths["generation"])
    make_error_figure(paths["errors"])
    return paths


def build_report() -> None:
    figures = make_figures()
    document = Document()
    configure_document(document)

    # 封面
    for _ in range(5):
        document.add_paragraph()
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(32)
    title_run = title_p.add_run(TITLE)
    set_run_font(title_run, east_asia="黑体", size=22, bold=True)
    type_p = document.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_run = type_p.add_run("项 目 报 告")
    set_run_font(type_run, east_asia="黑体", size=18, bold=True)
    for _ in range(8):
        document.add_paragraph()
    sub_p = document.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("视觉语言模型 · 页面级检索 · 多页视觉问答")
    set_run_font(sub_run, east_asia="宋体", size=13)
    document.add_page_break()

    # 中文摘要
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("摘  要")
    set_run_font(r, east_asia="黑体", size=16, bold=True)
    abstract = (
        "面向图表、表格、扫描页和复杂版式文档的问答任务，传统文本检索增强生成通常先执行光学字符识别，"
        "再对识别文本进行切分与向量检索。该链路容易丢失图表结构与页面布局信息，并将识别误差传递到后续检索和生成阶段。"
        "本项目设计并实现了一套以页面图像为检索单元的视觉语言模型检索增强问答系统。系统将文本 Query 与页面图像编码到可比较的表示空间，"
        "离线建立页面索引，在线召回 Top-K 证据页，并使用轻量级视觉语言模型完成逐页推理、检索分数加权融合或页面拼接问答。"
        "项目从 HuggingFaceM4/ChartQA 中整理出 2958 个去重页面和 6000 条问答样本，其中训练、验证和测试样本分别为 5000、500 和 500 条。"
        "检索实验比较了 SigLIP 零样本、PP-OCR+BGE、部分微调 SigLIP 和 ColSmol 四条链路。测试集结果表明，SigLIP 部分微调将 Recall@10 "
        "由 0.2520 提升至 0.3240，MRR@10 由 0.1903 提升至 0.2320；ColSmol 借助多向量晚交互取得最高的 Recall@10=0.4460 和 MRR@10=0.3669。"
        "在端到端问答实验中，ColSmol Top-3 逐页推理获得最高宽松准确率 0.1860；ColSmol Top-1 的宽松准确率为 0.1840，平均生成时延仅为 455.1 ms，"
        "表现出更合理的效果—成本折中。Oracle 证据页条件下宽松准确率为 0.3460，说明当前系统仍同时受到检索召回与轻量生成模型能力的限制。"
        "实验还显示，简单加入困难负样本并未进一步提高 MRR，而三页拼接虽然降低时延，却使宽松准确率下降至 0.1260。结果验证了页面级视觉检索的有效性，"
        "并说明端到端文档问答需要联合考虑检索表示、候选页组织方式与生成模型能力。"
    )
    add_body(document, abstract)
    add_body(document, "关键词：视觉语言模型；检索增强生成；页面图像检索；SigLIP；ColSmol；ChartQA", first_line=False, bold_prefix="关键词：")
    document.add_page_break()

    # English abstract
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ABSTRACT")
    set_run_font(r, east_asia="黑体", size=16, bold=True)
    en_abstract = (
        "Traditional retrieval-augmented generation pipelines for visually rich documents commonly rely on optical character recognition followed by text chunking and dense retrieval. "
        "Such pipelines may lose chart structure and page layout, while OCR errors propagate to retrieval and answer generation. This project designs and implements a page-image-based vision-language retrieval-augmented question answering system. "
        "Text queries and page images are encoded into comparable representations, page indexes are built offline, and Top-K evidence pages are retrieved online for visual question answering. "
        "A deduplicated subset of ChartQA containing 2,958 page images and 6,000 question-answer pairs is used, with 5,000/500/500 samples for training, development, and testing. "
        "Four retrieval pipelines are evaluated: zero-shot SigLIP, PP-OCR with BGE, partially fine-tuned SigLIP, and ColSmol. Partial fine-tuning improves test Recall@10 from 0.2520 to 0.3240 and MRR@10 from 0.1903 to 0.2320. "
        "ColSmol achieves the best retrieval performance with Recall@10 of 0.4460 and MRR@10 of 0.3669. In end-to-end evaluation, sequential reasoning over the Top-3 ColSmol pages yields the highest relaxed accuracy of 0.1860, whereas ColSmol Top-1 reaches a comparable 0.1840 at only 455.1 ms per query. "
        "The oracle-page relaxed accuracy is 0.3460, indicating that both retrieval and generation constrain the final result. Hard-negative variants do not improve MRR, and image collage reduces accuracy despite lower latency. These results demonstrate the feasibility of page-level visual RAG and clarify the trade-offs among representation granularity, evidence organization, accuracy, and inference cost."
    )
    add_body(document, en_abstract, first_line=False)
    add_body(document, "KEY WORDS: vision-language model; retrieval-augmented generation; page-image retrieval; SigLIP; ColSmol; ChartQA", first_line=False, bold_prefix="KEY WORDS:")
    document.add_page_break()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    set_run_font(r, east_asia="黑体", size=16, bold=True)
    add_toc(document)
    document.add_page_break()

    # 1 绪论
    add_heading(document, "1 绪论", 1)
    add_heading(document, "1.1 项目背景", 2)
    add_body(document, "企业文档中的有效信息不仅存在于连续文本中，还分布在表格行列、图表坐标、页面标题、脚注、印章以及多栏版式之中。传统文本 RAG 以 OCR 输出或 PDF 可提取文本为主要输入，通常要经历文字识别、版面解析、文本切片、向量化和检索等步骤。对于图表和复杂页面而言，文字序列不能完整表达视觉对象之间的空间关系；一旦 OCR 识别或阅读顺序出现错误，后续检索和生成阶段也会受到影响。")
    add_body(document, "RAG 的基本思想是将参数化生成模型与外部可检索知识结合，使回答可以依赖显式证据而不是仅依赖模型参数记忆[1]。稠密双编码检索器进一步证明，问题和文档可以分别编码为稠密向量，并通过内积或余弦相似度高效匹配[2]。在视觉文档场景中，本项目把检索单元从文本片段改为完整页面图像，以保留图表、字体和页面布局，并让系统返回可直接核验的证据页面。")

    add_heading(document, "1.2 问题定义", 2)
    add_body(document, "给定页面集合 P={p₁,p₂,…,pₙ}、自然语言问题 q 以及标准答案集合 A，本项目首先学习或调用跨模态检索函数 s(q,p)，从页面集合中选出得分最高的 K 个候选页；随后由视觉生成函数 g(q,p) 对候选页面进行回答，并通过融合函数 F 综合候选答案。系统目标可表示为：")
    add_equation(document, "P_K = TopK₍p∈P₎ s(q,p)，   â = F({g(q,p) | p∈P_K}, {s(q,p)})")
    add_body(document, "其中 P_K 为候选证据页集合，â 为最终预测答案。系统评价同时包含两个层面：一是正确证据页能否进入 Top-K，二是在给定候选页后能否输出正确答案。将两类指标分开，可以定位错误究竟来自检索阶段还是生成阶段。")

    add_heading(document, "1.3 项目内容与主要工作", 2)
    add_numbered(document, [
        "建立统一的页面级数据模式，完成 ChartQA 页面、Query、答案和证据页的接入、去重、划分与完整性验证。",
        "实现 SigLIP 单向量视觉检索、PP-OCR+BGE 文本检索和 ColSmol 多向量晚交互检索，形成可比较的三类预训练基线。",
        "基于 Query—证据页监督信号，对 SigLIP 文本塔和视觉塔末端层进行部分微调，使用多正例 InfoNCE 目标，并比较普通批内负例与困难负例策略。",
        "实现 Top-1、Top-3 逐页推理、检索分数加权投票和 Top-3 页面拼接问答，在同一测试集上评估检索召回、EM、宽松准确率与生成时延。",
        "形成可复现的脚本、数据清单、模型检查点、逐条预测与汇总报告，使实验结果能够追溯到具体 Query 和证据页面。",
    ])

    add_heading(document, "1.4 阶段划分", 2)
    add_table(document, ["阶段", "核心任务", "主要产出"], [
        ["阶段一", "数据接入与最小闭环", "统一模式、ChartQA 子集、页面索引与最小检索验证"],
        ["阶段二", "补齐对照基线", "SigLIP、OCR+BGE、视觉晚交互和页面问答对照链路"],
        ["阶段三", "检索器训练与消融", "SigLIP 部分微调、困难负样本实验、检索指标表"],
        ["阶段四", "检索—生成联合评估", "Top-1/Top-3、逐页/拼接、多页答案融合与误差分解"],
    ], widths=[2.4, 5.5, 8.0])

    # 2 理论基础
    add_heading(document, "2 理论基础与关键技术", 1)
    add_heading(document, "2.1 检索增强生成与页面级 RAG", 2)
    add_body(document, "标准 RAG 通常由检索器、外部索引和生成器三部分组成[1]。检索器根据 Query 选择相关上下文，生成器在这些上下文条件下回答问题。页面级 RAG 沿用这一结构，但知识单元不再是纯文本 chunk，而是包含完整视觉布局的页面图像。页面向量可离线计算，在线仅需计算 Query 表示并完成相似度检索，因此仍保留稠密检索的工程优势。")
    add_body(document, "与直接把全部文档交给大模型相比，先检索后生成能缩小视觉上下文、减少无关页面干扰，并保留证据来源。对 ChartQA 而言，一个页面就是一张图表；在真实 PDF 中，一个页面还可能包含段落、表格和签章。项目统一使用 page_id 标识证据，从而使检索结果、答案和评测记录能够关联。")

    add_heading(document, "2.2 VLM、Embedding 与跨模态双塔", 2)
    add_body(document, "视觉语言模型能够联合处理图像与文本。检索场景下，模型不一定直接生成自然语言，而是输出能够表示语义的向量或多向量。SigLIP 采用图像编码器和文本编码器学习图文对齐表示，其预训练目标使用逐对 sigmoid 损失，减少了对全局 softmax 归一化的依赖[3]。本项目使用 google/siglip-base-patch16-224，将页面图像与 Query 分别编码并进行 L2 归一化。")
    add_equation(document, "q = f_text(Query) / ‖f_text(Query)‖₂，   vᵢ = f_image(Pageᵢ) / ‖f_image(Pageᵢ)‖₂")
    add_equation(document, "s(q,vᵢ) = qᵀvᵢ = cos(q,vᵢ)")
    add_body(document, "由于页面向量在查询前已经写入索引，在线流程只需要计算一次问题向量。当前实现使用精确遍历计算余弦相似度，适用于 2958 页的实验规模；它与 FAISS 等近似最近邻索引在逻辑上属于同一检索阶段，但本报告不把未实际使用的外部向量数据库写入实现结果。")

    add_heading(document, "2.3 InfoNCE 对比学习", 2)
    add_body(document, "部分微调采用多正例 InfoNCE。一个批次内，每个 Query 对应至少一个正例页面，其他页面作为负例；当相同页面对应多个问题时，正例掩码允许多个匹配关系。对第 i 个 Query，损失可写为：")
    add_equation(document, "Lᵢ = −log [ Σⱼ∈P(i) exp(sᵢⱼ/τ) / Σₖ exp(sᵢₖ/τ) ]")
    add_body(document, "其中 P(i) 是该问题的正例页集合，τ 是温度系数。项目设置 τ=0.07。InfoNCE 通过提高正例相似度、降低负例相似度，使 Query 与证据页在同一表示空间中更容易分离。对比预测编码工作提出了 InfoNCE 形式的信息估计目标[9]，本项目将其用于有监督图文检索。")

    add_heading(document, "2.4 OCR+BGE 文本基线", 2)
    add_body(document, "OCR 基线先使用 PP-OCRv6-medium 从页面图片中提取文字，过滤置信度低于 0.5 的文本行，再将同页文字拼接为一个页面文档。之后使用 BAAI/bge-small-en-v1.5 分别编码 Query 与 OCR 文本并进行余弦检索。BGE 属于通用文本嵌入模型，相关工作通过多任务数据与统一训练方法提升了文本表示能力[6]。该链路是传统文本 RAG 的典型对照：它可以利用清晰文字，但图表坐标关系和视觉布局会被压缩为线性字符串。")

    add_heading(document, "2.5 ColSmol 多向量晚交互", 2)
    add_body(document, "单向量方法把整页压缩成一个向量，局部文字、图例和数据点可能在池化过程中被弱化。ColPali 将页面图像直接编码成多向量，并通过晚交互匹配 Query token 与页面视觉 token，在视觉丰富文档检索中表现突出[4]。ColBERTv2 也说明了多向量晚交互相对于单向量表示的细粒度匹配优势[5]。本项目使用 vidore/colSmol-500M，通过 ColIdefics3Processor 的 score_multi_vector 计算 Query 与页面的 MaxSim 风格得分。")

    add_heading(document, "2.6 视觉问答与多页融合", 2)
    add_body(document, "生成模块采用 HuggingFaceTB/SmolVLM-500M-Instruct。SmolVLM 面向低资源视觉理解，500M 版本在较小参数规模下兼顾文档问答和图表理解[11]。项目对每个候选页使用相同提示词，要求模型仅依据图表图片回答，并仅返回简短答案；解码采用确定性生成，最大新增长度为 16 token。")
    add_body(document, "Top-3 逐页方案分别对三页推理，再对检索得分执行 softmax 归一化并按规范化答案聚合权重。若不同页面给出相同答案，其权重相加；最终选择总权重最高的答案。拼接方案则把三页纵向缩放并合成一张图，仅调用一次生成器。二者分别代表较高视觉分辨率但多次推理的方案，以及较低成本但页面内容被压缩的方案。")

    add_heading(document, "2.7 评价指标", 2)
    add_table(document, ["指标", "定义", "解释"], [
        ["Recall@K", "正确证据页进入前 K 名的问题比例", "衡量检索覆盖率，K 取 1、3、10"],
        ["MRR@10", "前 10 名中首个正确证据页倒数排名的均值", "正确页越靠前，指标越高"],
        ["EM", "规范化后预测答案与任一标准答案完全一致", "严格衡量答案一致性"],
        ["Relaxed Accuracy", "EM 或数值答案在目标值 ±5% 范围内", "适配 ChartQA 数值回答口径"],
        ["Accuracy given hit", "已召回正确证据页样本中的宽松准确率", "隔离观察生成器表现"],
        ["Mean ms", "每个方法的平均生成时延", "评估效果与推理成本"],
    ], widths=[3.2, 7.4, 5.4])

    # 3 需求与架构
    add_heading(document, "3 需求分析与总体设计", 1)
    add_heading(document, "3.1 功能需求", 2)
    add_bullets(document, [
        "接收页面图像、Query、标准答案与证据页标注，并验证字段和文件路径的完整性。",
        "支持离线编码页面、持久化索引、在线检索 Top-K 页面并返回页面编号、排名和得分。",
        "支持零样本视觉检索、OCR 文本检索、视觉晚交互检索以及微调检索器之间的统一评测。",
        "支持 Oracle、Top-1、Top-3 逐页推理和 Top-3 拼接推理，并保存逐条答案以便断点续跑。",
        "输出 JSON、JSONL、CSV、Markdown、模型检查点和日志，使训练过程与结果可复现。",
    ])

    add_heading(document, "3.2 非功能需求", 2)
    add_body(document, "系统强调可追溯性、可复现性和模块可替换性。页面与问题使用稳定 ID 关联；数据划分、模型名称、Top-K、学习率和随机种子均进入配置或结果文件；页面特征和问答结果采用缓存，避免重复计算；训练、检索和生成分别由独立脚本执行，便于在本地 CPU 环境、RTX 5090 和 RTX 4090 云实例间迁移。")

    add_heading(document, "3.3 总体架构与数据流", 2)
    add_figure(document, figures["architecture"], "图 3-1 页面级 VLM-RAG 总体架构与数据流", width=6.35)
    add_body(document, "离线链路负责把页面图像转换为可检索表示。ChartQA 适配器读取原始图表图像，按 PNG 内容哈希去重，生成 pages.jsonl、samples.jsonl 和 manifest.json；随后不同检索器分别建立单向量、文本向量或多向量页面索引。在线链路接收 Query，编码后从索引中返回 Top-K 页面；生成器读取页面并输出候选答案，最后由融合模块选出答案并保留证据页。")

    add_heading(document, "3.4 系统阶段间的数据对象", 2)
    add_table(document, ["数据对象", "主要字段", "作用"], [
        ["DocumentPage", "page_id、doc_id、doc_type、page_no、image_path、metadata", "描述页面图像及来源信息"],
        ["DocumentQASample", "query_id、query、answers、evidence_page_ids、split", "描述问题、标准答案和证据监督"],
        ["页面索引", "page_ids、embedding、模型与维度元数据", "保存离线编码结果"],
        ["Retrieval Result", "query_id、top_10、first_evidence_rank", "记录每条问题的候选页和得分"],
        ["Generation Prediction", "method、page_ids、prediction、error_type、generation_ms", "记录端到端答案、证据和错误类型"],
    ], widths=[3.5, 7.7, 4.8])

    # 4 数据与实现
    add_heading(document, "4 数据构建与系统实现", 1)
    add_heading(document, "4.1 ChartQA 数据接入", 2)
    add_body(document, "ChartQA 是面向图表视觉与逻辑推理的问答基准，原始数据包含人工问题与基于图表摘要生成的问题[7]。本项目通过 HuggingFaceM4/ChartQA 的流式接口分别读取 train、val 和 test，其中 val 映射为 dev。每张图片统一保存为 PNG，并对图片字节计算 SHA-256；相同图片只保留一个页面对象。若同一图片出现在不同源划分，适配器仅保留首次出现的划分，防止相同页面跨训练、验证和测试集合泄漏。")
    add_table(document, ["数据项", "规模/取值"], [
        ["去重页面图像", "2958"],
        ["问答样本", "6000"],
        ["训练集", "5000"],
        ["验证集", "500"],
        ["测试集", "500"],
        ["文档类型", "chart"],
        ["证据粒度", "单个图表页面"],
    ], widths=[7.0, 8.0])
    add_body(document, "阶段一的核心产物不是模型权重，而是后续所有实验共享的标准数据闭环：pages/ 保存图像，pages.jsonl 保存页面元数据，samples.jsonl 保存 Query、答案和证据页，manifest.json 汇总规模、划分和警告。validate_bundle 会检查重复 ID、空问题、空答案、缺失证据页、不支持的 split 和潜在文档泄漏。")

    add_heading(document, "4.2 核心模块实现", 2)
    add_table(document, ["代码文件", "功能说明"], [
        ["dataset_adapters.py", "流式读取 ChartQA，统一图片格式，按哈希去重并构造页面与问答对象。"],
        ["dataset_schema.py", "定义 DocumentPage、DocumentQASample、DatasetBundle，并完成保存、加载与校验。"],
        ["siglip_encoder.py", "调用 Transformers 加载 SigLIP，批量编码 Query 和页面并进行向量归一化。"],
        ["siglip_index.py", "保存页面单向量和元数据，按余弦相似度执行精确 Top-K 检索。"],
        ["ocr_extractor.py", "封装 PaddleOCR，过滤低置信度文本行并记录单页 OCR 时延。"],
        ["text_encoder.py", "调用 SentenceTransformer/BGE 编码 OCR 文本和 Query。"],
        ["colsmol_encoder.py", "加载 ColIdefics3/ColSmol，生成页面与问题多向量并计算晚交互得分。"],
        ["retrieval_adapter.py", "实现残差低秩投影头，为冻结特征上的轻量适配实验提供双塔或共享塔结构。"],
        ["smolvlm_generator.py", "构造图文提示词，调用 SmolVLM 确定性生成简短答案。"],
        ["train_siglip_partial_finetune.py", "冻结大部分 SigLIP，仅解冻两塔末端层，以多正例 InfoNCE 训练并早停。"],
        ["evaluate_stage4_generation.py", "汇总检索结果，缓存逐页答案，执行 Top-1/Top-3/拼接策略并计算端到端指标。"],
    ], widths=[5.4, 10.2], font_size=10)

    add_heading(document, "4.3 训练与评估脚本", 2)
    add_body(document, "scripts 目录承担可执行入口的职责。download_chartqa.py 构建真实子集；validate_training_dataset.py 输出校验报告；evaluate_siglip_retrieval.py 建立零样本 SigLIP 索引；extract_chartqa_ocr.py 与 evaluate_ocr_retrieval.py 完成 OCR+BGE；evaluate_colsmol_retrieval.py 完成多向量检索；train_siglip_partial_finetune.py 负责部分微调；evaluate_siglip_partial_checkpoint.py 重新编码全部页面与 Query 并评估检查点；summarize_stage3_experiments.py 和 summarize_large_retrieval_baselines.py 生成对照表；run_stage4_generation_suite.sh 串联最终生成实验。")
    add_body(document, "每一步的输出目录互相独立，脚本不会覆盖不同方法的结果。阶段四把 page_answers.jsonl 和 collage_answers.jsonl 作为缓存，键分别为 (query_id,page_id) 和 (query_id,retriever)，因此网络或终端中断后可以从剩余样本继续执行。run_config.json 会在续跑前校验模型、Top-K 和检索结果路径，防止混用不同实验配置。")

    add_heading(document, "4.4 页面检索与问答流程", 2)
    add_numbered(document, [
        "离线读取全部 DocumentPage，使用指定模型批量编码页面并写入索引。",
        "对每条 DocumentQASample 编码 Query，与全部页面表示计算相似度，保存前 10 个候选。",
        "根据 evidence_page_ids 找到首个证据页排名，计算 Recall@1、Recall@3、Recall@10 和 MRR@10。",
        "阶段四从测试集读取两种检索器的 Top-3；对 Oracle 页及候选页分别调用 SmolVLM。",
        "Top-1 直接采用第一候选页答案；Top-3 逐页方案按检索分数加权投票；拼接方案合成一张图后生成一次。",
        "将预测答案与标准答案规范化比较，输出 EM、宽松准确率、条件准确率、时延和错误类型。",
    ])

    add_heading(document, "4.5 实验实现边界", 2)
    add_body(document, "本报告中的全部数值均来自真实模型运行结果，但实验语料只覆盖 ChartQA 图表类型。ChartQA 的每个 page_id 实际上是一张独立图表，不是同一企业 PDF 中连续的多个页面。因此阶段四的“Top-3 多页”准确含义是从全库取三个候选页面分别推理或拼接比较，而不是对同一文档内跨页事实进行联合推理。报告不将该实验外推为对合同、PPT、制度等六类文档的已验证结论。")

    # 5 训练与基线
    add_heading(document, "5 检索器训练与基线设计", 1)
    add_heading(document, "5.1 实验环境", 2)
    add_table(document, ["项目", "配置"], [
        ["云端训练 GPU", "NVIDIA GeForce RTX 5090 32GB（SigLIP 部分微调）"],
        ["补充评估 GPU", "NVIDIA GeForce RTX 4090 24GB（OCR/ColSmol/Stage 4）"],
        ["系统镜像", "Ubuntu 22.04，Python 3.12，PyTorch 2.8.0，CUDA 12.8 镜像"],
        ["核心框架", "Transformers、PaddleOCR、SentenceTransformers、colpali-engine"],
        ["数据规模", "2958 页，6000 Query；训练/验证/测试=5000/500/500"],
        ["随机种子", "42"],
    ], widths=[4.6, 10.6])

    add_heading(document, "5.2 SigLIP 零样本基线", 2)
    add_body(document, "零样本基线直接加载 google/siglip-base-patch16-224，不更新参数。页面和 Query 通过 get_image_features 与 get_text_features 编码为归一化单向量，再对全部页面执行余弦排序。该基线是同一模型架构下最基础的对照，可用于判断后续微调是否真正改善检索，而不受不同模型容量和评分方式影响。")

    add_heading(document, "5.3 SigLIP 双塔部分微调", 2)
    add_body(document, "训练时先冻结全部参数，再解冻文本编码器最后两层、视觉编码器最后两层，以及两塔的归一化和投影头。总参数量为 203,155,970，可训练参数为 36,032,256，占 17.74%。这种方式保留预训练图文对齐能力，同时让高层语义适配 ChartQA 的 Query—页面关系，并将显存与训练成本控制在单张消费级 GPU 可承受范围。")
    add_table(document, ["超参数", "取值"], [
        ["Epoch", "4"], ["Batch size", "16"], ["梯度累积", "2"], ["评估 Batch", "32"],
        ["学习率", "2×10⁻⁶"], ["权重衰减", "0.01"], ["InfoNCE 温度", "0.07"],
        ["Warmup 比例", "0.05"], ["Early stopping patience", "2"], ["解冻层数", "文本 2 层 + 视觉 2 层"],
    ], widths=[7.5, 7.5])
    add_body(document, "优化器采用 AdamW，学习率经过线性 warmup 后线性衰减；使用 bfloat16 自动混合精度、梯度累积与 1.0 的梯度裁剪。每个 Epoch 后在验证集上重新编码页面和 Query，以 MRR@10 选择最佳检查点。最终保存的 siglip_partial_finetune.pt 只包含可训练层状态、基础模型名称、可训练参数名称与配置，便于在基础模型上恢复。")

    add_heading(document, "5.4 困难负样本消融", 2)
    add_body(document, "普通训练把批内其他页面视为负例。困难负样本实验从 SigLIP 零样本检索结果中选择排名靠前但不属于 evidence_page_ids 的页面，并分别为每个 Query 加入 1 个或 2 个负例。其目的在于让模型区分“视觉或语义相似但不是答案页”的候选。然而，ChartQA 中不同 Query 可能共享相似图表，且证据标注只指出当前问题对应页面，靠前的非证据页面可能包含未标注但语义相近的内容。若直接视为严格负例，会引入假负例。")

    add_heading(document, "5.5 三类检索基线的作用", 2)
    add_table(document, ["方法", "表示方式", "在实验中的作用"], [
        ["SigLIP 零样本", "Query 单向量 + 页面单向量", "最基础同架构对照，衡量预训练模型直接迁移能力"],
        ["PP-OCR+BGE", "Query 文本向量 + OCR 页面文本向量", "传统文本 RAG 对照，观察 OCR 链路的效果"],
        ["ColSmol", "Query 多向量 + 页面视觉多向量晚交互", "强视觉检索基线，衡量局部 token 级匹配价值"],
        ["SigLIP 部分微调", "经过任务适配的双塔单向量", "本项目实际训练方法，验证监督微调增益"],
    ], widths=[3.7, 5.4, 6.7])

    # 6 实验
    add_heading(document, "6 实验结果与分析", 1)
    add_heading(document, "6.1 SigLIP 训练过程", 2)
    add_figure(document, figures["training"], "图 6-1 SigLIP 部分微调的训练损失与验证集 MRR@10", width=5.8)
    add_body(document, "训练损失从第 1 轮的 2.0723 下降至第 4 轮的 1.3690，下降约 33.9%，说明优化目标得到持续拟合。验证集 MRR@10 从第 1 轮的 0.2511 上升至第 3 轮的 0.2591，第 4 轮轻微回落至 0.2588，因此最佳检查点选择第 3 轮。与微调前验证集 MRR@10=0.2123 相比，最佳值提高约 22.0%。")
    add_table(document, ["Epoch", "Train Loss", "Dev R@1", "Dev R@3", "Dev R@10", "Dev MRR@10"], [
        ["1", "2.0723", "0.2100", "0.2780", "0.3440", "0.2511"],
        ["2", "1.6189", "0.2100", "0.2840", "0.3500", "0.2551"],
        ["3", "1.4352", "0.2140", "0.2860", "0.3580", "0.2591"],
        ["4", "1.3690", "0.2140", "0.2840", "0.3580", "0.2588"],
    ], widths=[2.0, 3.1, 2.5, 2.5, 2.7, 3.1], font_size=10)

    add_heading(document, "6.2 检索基线对比", 2)
    add_table(document, ["排名", "方法", "Recall@1", "Recall@3", "Recall@10", "MRR@10"], [
        ["1", "ColSmol", "0.3320", "0.3900", "0.4460", "0.3669"],
        ["2", "SigLIP 部分微调", "0.1900", "0.2600", "0.3240", "0.2320"],
        ["3", "PP-OCR+BGE", "0.1980", "0.2540", "0.2800", "0.2274"],
        ["4", "SigLIP 零样本", "0.1600", "0.2160", "0.2520", "0.1903"],
    ], widths=[1.6, 5.1, 2.4, 2.4, 2.7, 2.7], font_size=10)
    add_figure(document, figures["retrieval"], "图 6-2 四种检索方法在测试集上的指标对比", width=6.1)
    add_body(document, "与零样本 SigLIP 相比，部分微调使 Recall@1、Recall@3、Recall@10 和 MRR@10 分别提高 0.0300、0.0440、0.0720 和 0.0416，相对增幅分别为 18.75%、20.37%、28.57% 和 21.88%。这表明 5000 条训练 Query 提供的监督信号能够有效调整高层图文表示。")
    add_body(document, "OCR+BGE 的 Recall@1=0.1980，略高于微调 SigLIP 的 0.1900；但其 Recall@10=0.2800 和 MRR@10=0.2274 均低于微调模型。这说明 OCR 文本在部分问题上可以把正确页直接排到首位，但整体候选覆盖和平均排序质量稍弱。OCR 日志还显示 2958 页均成功获得文本，单页平均处理时间约 7571.2 ms；该时间属于离线 OCR 处理，不应与阶段四的单次生成时延直接比较。")
    add_body(document, "ColSmol 在所有检索指标上均为最高，其测试集 MRR@10 比微调 SigLIP 高 0.1349，Recall@10 高 0.1220。结果支持多向量晚交互对图表局部元素匹配的优势：Query 中的实体、年份或数值关系可分别与页面不同视觉 token 匹配，而不必把整页信息压缩为单个向量。但 ColSmol 的页面编码耗时为 1168.9 s，峰值 GPU 显存约 5605 MB，说明其更高检索质量伴随更复杂的离线表示与存储。")

    add_heading(document, "6.3 困难负样本结果", 2)
    add_table(document, ["方法", "Recall@1", "Recall@3", "Recall@10", "MRR@10"], [
        ["SigLIP 零样本", "0.1600", "0.2160", "0.2520", "0.1903"],
        ["部分微调", "0.1900", "0.2600", "0.3240", "0.2320"],
        ["部分微调 + 1 个困难负例", "0.1840", "0.2600", "0.3240", "0.2313"],
        ["部分微调 + 2 个困难负例", "0.1660", "0.2640", "0.3240", "0.2215"],
    ], widths=[6.4, 2.4, 2.4, 2.7, 2.7], font_size=10)
    add_body(document, "加入 1 个困难负例后 MRR@10 从 0.2320 轻微降至 0.2313；加入 2 个后降至 0.2215。两个变体的 Recall@10 都保持 0.3240，第二个变体的 Recall@3 略升至 0.2640，但 Recall@1 明显下降。这意味着困难负例没有提升整体排序质量，额外约束主要改变了前几名内部排序。")
    add_body(document, "该结果只能说明当前静态挖掘策略未带来收益，不能推导出困难负例方法本身无效。一方面，负例来自微调前 SigLIP 排名，可能含有视觉相似或可回答同类问题的假负例；另一方面，为容纳额外页面，三个实验的单步 batch size 分别为 16、12 和 8，尽管使用梯度累积保持接近的有效批量，训练构成仍非完全相同。因此本消融结论应限定在当前数据和实现条件下。")

    add_heading(document, "6.4 端到端生成结果", 2)
    add_table(document, ["方法", "检索召回", "EM", "宽松准确率", "命中条件准确率", "平均时延/ms"], [
        ["Oracle 证据页", "1.0000", "0.3180", "0.3460", "0.3460", "466.3"],
        ["SigLIP 微调 Top-1", "0.1900", "0.1280", "0.1620", "0.4632", "465.0"],
        ["SigLIP 微调 Top-3 逐页", "0.2600", "0.1280", "0.1640", "0.3462", "1387.0"],
        ["ColSmol Top-1", "0.3320", "0.1540", "0.1840", "0.3675", "455.1"],
        ["ColSmol Top-3 逐页", "0.3900", "0.1580", "0.1860", "0.3333", "1369.5"],
        ["ColSmol Top-3 拼接", "0.3900", "0.0980", "0.1260", "0.2000", "350.3"],
    ], widths=[5.1, 2.5, 1.8, 2.7, 3.1, 2.8], font_size=9.5)
    add_figure(document, figures["generation"], "图 6-3 不同检索—生成组合的宽松准确率与平均时延", width=6.25)
    add_body(document, "Oracle 方案直接把标准证据页交给 SmolVLM，在 500 条测试问题上得到 159 条 EM 正确和 173 条宽松正确，宽松准确率为 0.3460。由于检索召回固定为 1.0，这一结果反映当前生成模型、提示词和图表推理能力共同形成的上限。即使页面完全正确，仍有 327 条宽松错误，因此不能把所有端到端失败归因于检索。")
    add_body(document, "微调 SigLIP 从 Top-1 扩展到 Top-3 后，检索召回由 0.1900 提高至 0.2600，但宽松准确率仅从 0.1620 提高至 0.1640，平均生成时延从 465.0 ms 增长到 1387.0 ms，接近三倍。ColSmol 也呈现相同趋势：Top-3 逐页召回从 0.3320 提升至 0.3900，宽松准确率只提高 0.0020，而时延从 455.1 ms 增至 1369.5 ms。结果说明当前基于检索分数的简单答案投票没有充分转化新增证据页的召回收益。")
    add_body(document, "ColSmol Top-3 逐页方案以 EM=0.1580、宽松准确率=0.1860取得最高端到端结果；然而 ColSmol Top-1 的宽松准确率已达到 0.1840，差距仅为 0.0020，时延约为前者的三分之一。因此在当前模型与数据上，Top-1 是更有吸引力的效果—成本折中。")
    add_body(document, "Top-3 拼接保持 0.3900 的证据召回，但宽松准确率下降至 0.1260，比逐页方案低 0.0600；其平均时延为 350.3 ms，比逐页方案低约 74.4%。拼接把三张图表压缩进固定画布，虽然只生成一次，却降低了文本、刻度和数据点的有效分辨率，并增加视觉干扰。该对照证明，多候选页面的组织方式会显著影响生成结果，不能仅根据检索召回判断端到端性能。")

    add_heading(document, "6.5 错误来源分析", 2)
    add_figure(document, figures["errors"], "图 6-4 测试集端到端错误来源分解", width=6.0)
    add_body(document, "阶段四实现按“回答正确优先”的方式标记错误：若宽松匹配正确，则记为 correct；若答案错误但证据页已进入候选，记为 generation_error；若答案错误且证据页未进入候选，记为 retrieval_miss。因此图 6-4 中的检索未命中数量不等于 500×(1−Recall)，因为个别非标注页面也可能偶然生成正确答案。")
    add_body(document, "SigLIP Top-1 有 81 条宽松正确、368 条“未命中且回答错误”和 51 条“已命中但生成错误”；ColSmol Top-1 的对应数量为 92、303 和 105。ColSmol 通过减少检索缺失提高了最终正确数，但也暴露出更多已命中后的生成错误。Top-3 逐页方案进一步减少检索缺失，却增加生成错误，表明候选答案之间可能发生冲突，而检索分数并不是答案置信度的充分替代。")
    add_body(document, "从系统角度看，当前瓶颈具有层级性：SigLIP 单向量方法首先受到页面召回限制；ColSmol 提高召回后，SmolVLM 的图表读取与计算能力成为更明显的限制；Top-3 方案增加候选后，简单加权投票又成为证据融合限制。Oracle、检索器与生成策略三组实验共同构成了这一判断的证据链。")

    add_heading(document, "6.6 结果有效性讨论", 2)
    add_body(document, "所有检索方法使用同一 500 条测试集和同一 evidence_page_ids 标注，阶段四也复用阶段三输出的逐条 Top-10 排名，因此同表内部具备可比性。微调模型通过验证集 MRR 选取检查点，测试集只用于最终报告。结果文件保留完整 Query、候选页、排名、预测答案和耗时，便于重新统计。")
    add_body(document, "实验仍存在明确的解释边界。第一，数据只包含图表，不能代表合同、报表、PPT、单据、手册和制度等多业务文档；第二，页面池规模为 2958，当前精确搜索结果不能直接推断百万页索引性能；第三，ColSmol 与 SigLIP 的表示粒度、参数量和评分方式不同，ColSmol 的领先属于系统方法对比而非同容量消融；第四，阶段四生成器使用 500M 轻量模型与固定短答案提示，Oracle 结果表明生成上限较低。上述边界不会否定当前结果，但限定了结论适用范围。")

    # 7 结论
    add_heading(document, "7 结论", 1)
    add_body(document, "本项目完成了从真实公开数据接入、页面级索引、跨模态检索器微调、传统 OCR 基线、视觉晚交互基线到多候选页面问答的完整实验闭环。系统以页面图像作为 RAG 检索单元，将 Query、答案和证据页作为核心监督与评测对象，并通过标准化 JSONL 模式、独立脚本和结果缓存保证实验可追溯。")
    add_body(document, "检索实验表明，SigLIP 部分微调相对于零样本模型显著改善全部测试指标，验证了使用领域问答对进行双塔适配的价值；困难负例的当前构造方式没有继续提高 MRR，说明负样本质量比数量更关键。三类基线中，ColSmol 多向量晚交互获得最好的页面召回与排序质量，说明图表检索受益于细粒度视觉匹配。")
    add_body(document, "端到端实验进一步表明，更高的 Top-K 召回不会自动转化为同等幅度的答案提升。ColSmol Top-3 逐页推理取得最高宽松准确率 0.1860，但 ColSmol Top-1 以 0.1840 的相近效果和约三分之一时延形成更好的工程折中；直接拼接三页虽更快，却明显损失准确率。Oracle 条件下仍只有 0.3460 的宽松准确率，证明检索器和生成器都需要纳入系统级评价。总体而言，页面级视觉 RAG 在复杂图表问答中具有可行性，而检索表示粒度、候选页组织方式和生成能力共同决定最终性能。")

    # 参考文献
    add_heading(document, "参考文献", 1)
    references = [
        "[1] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems. 2020, 33: 9459-9474.",
        "[2] KARPUKHIN V, OĞUZ B, MIN S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]//Proceedings of EMNLP. 2020: 6769-6781.",
        "[3] ZHAI X, MUSTAFA B, KOLESNIKOV A, et al. Sigmoid Loss for Language Image Pre-Training[C]//Proceedings of ICCV. 2023: 11975-11986.",
        "[4] FAYSSE M, SIBILLE H, WU T, et al. ColPali: Efficient Document Retrieval with Vision Language Models[EB/OL]. arXiv:2407.01449, 2024.",
        "[5] SANTHANAM K, KHATTAB O, SAAD-FALCON J, et al. ColBERTv2: Effective and Efficient Retrieval via Lightweight Late Interaction[C]//Proceedings of NAACL. 2022: 3715-3734.",
        "[6] XIAO S, LIU Z, ZHANG P, et al. C-Pack: Packed Resources for General Chinese Embeddings[EB/OL]. arXiv:2309.07597, 2023.",
        "[7] MASRY A, LONG D X, TAN J Q, et al. ChartQA: A Benchmark for Question Answering about Charts with Visual and Logical Reasoning[C]//Findings of ACL. 2022: 2263-2279.",
        "[8] DU Y, LI C, GUO R, et al. PP-OCRv2: Bag of Tricks for Ultra Lightweight OCR System[EB/OL]. arXiv:2109.03144, 2021.",
        "[9] VAN DEN OORD A, LI Y, VINYALS O. Representation Learning with Contrastive Predictive Coding[EB/OL]. arXiv:1807.03748, 2018.",
        "[10] DOSOVITSKIY A, BEYER L, KOLESNIKOV A, et al. An Image is Worth 16×16 Words: Transformers for Image Recognition at Scale[C]//Proceedings of ICLR. 2021.",
        "[11] MARAFIOTI A, ZOHAR O, FARRÉ M, et al. SmolVLM: Redefining Small and Efficient Multimodal Models[EB/OL]. arXiv:2504.05299, 2025.",
        "[12] VIDORE. colSmol-500M Model Card[EB/OL]. https://huggingface.co/vidore/colSmol-500M, 2026-08-08.",
        "[13] HUGGING FACE. SmolVLM-500M-Instruct Model Card[EB/OL]. https://huggingface.co/HuggingFaceTB/SmolVLM-500M-Instruct, 2026-08-08.",
        "[14] VASWANI A, SHAZEER N, PARMAR N, et al. Attention Is All You Need[C]//Advances in Neural Information Processing Systems. 2017, 30.",
    ]
    for ref in references:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(ref)
        set_run_font(run, size=10.5)

    # 附录
    add_heading(document, "附录 A 主要实验产物", 1)
    add_table(document, ["路径", "内容"], [
        ["data/real/chartqa_large/manifest.json", "数据规模、划分和校验警告"],
        ["outputs/siglip_chartqa_large_baseline/", "SigLIP 零样本检索指标与逐条 Top-10"],
        ["outputs/siglip_partial_full/", "部分微调检查点、训练日志与指标"],
        ["outputs/siglip_partial_hardneg_1/", "每条 Query 加 1 个困难负例的结果"],
        ["outputs/siglip_partial_hardneg_2/", "每条 Query 加 2 个困难负例的结果"],
        ["outputs/ocr_bge_chartqa_large/", "PP-OCR+BGE 检索结果及 OCR 时延"],
        ["outputs/colsmol_chartqa_large/", "ColSmol 多向量检索结果及资源统计"],
        ["outputs/stage4_generation_test/", "生成配置、缓存、逐条预测、指标与对照表"],
    ], widths=[8.0, 7.2], font_size=10)

    add_headers_and_footers(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
