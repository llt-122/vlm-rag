from __future__ import annotations

import csv
import json
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCS = PROJECT_ROOT / "docs"
SOURCE = DOCS / "VLM-RAG项目报告_排版放大修订版.docx"
OUTPUT = DOCS / "VLM-RAG项目报告_融合多页验证版.docx"
DELIVERY = PROJECT_ROOT / "outputs" / "phase2_delivery"
MULTIPAGE = PROJECT_ROOT / "data" / "multipage" / "finance_40"
MULTI_OUTPUT = PROJECT_ROOT / "outputs" / "multipage_colsmol_top3"

HEADING1_PPR = None
HEADING2_PPR = None
BODY_PPR = None
CAPTION_PPR = None
TABLE_PR = None
TABLE_HEADER_TRPR = None
TABLE_BODY_TRPR = None
TABLE_HEADER_TCPR = None
TABLE_BODY_TCPR = None
TABLE_HEADER_PPR = None
TABLE_BODY_PPR = None


def main() -> None:
    doc = Document(SOURCE)
    _capture_format_templates(doc)
    conclusion = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.style.name == "Heading 1" and paragraph.text.strip().startswith("8 ")
    )
    conclusion.text = "9 结论"
    _format_heading_runs(conclusion, 18)

    body = doc.element.body
    old_last = next(
        element for element in reversed(list(body.iterchildren()))
        if element.tag != qn("w:sectPr")
    )
    _append_phase2(doc)
    all_elements = list(body.iterchildren())
    anchor_index = all_elements.index(old_last)
    new_elements = [
        element for element in all_elements[anchor_index + 1:]
        if element.tag != qn("w:sectPr")
    ]
    for element in new_elements:
        conclusion._p.addprevious(element)

    doc.save(OUTPUT)
    print(OUTPUT)


def _append_phase2(doc: Document) -> None:
    retrieval = _csv(DELIVERY / "five_retrieval_baselines.csv")
    generation = _csv(DELIVERY / "four_smolvlm_qa_schemes.csv")
    failure_summary = _json(DELIVERY / "retrieval_failure_summary.json")
    multipage_manifest = _json(MULTIPAGE / "manifest.json")
    metrics = _json(MULTI_OUTPUT / "metrics.json")
    predictions = _json(MULTI_OUTPUT / "predictions.json")

    _heading(doc, "8 多页连续文档验证与二阶段实验", 1)
    _body(doc, (
        "前述 ChartQA 实验以单张图表页面为一个独立样本，Top-3表示从整个页面库中返回三个候选页面，"
        "并不等价于同一份文档内部的跨页联合推理。为检验系统对连续多页数据的适应能力，本阶段首先"
        "统一复核检索与问答基线，然后构建具有文档归属、连续页码及多证据标注的小型验证集，实际运行"
        "ColSmol检索与SmolVLM多图问答。既有指标来自真实实验产物；新多页数据属于程序化构造的受控"
        "验证集，不作为真实企业数据使用。"
    ))

    _heading(doc, "8.1 五种检索方案的统一基线", 2)
    _body(doc, (
        "五种方案统一采用ChartQA测试集500条Query进行比较。SigLIP难负例方案取每个Query加入1个"
        "难负例的实验；每个Query加入2个难负例的结果作为补充消融，不重复进入主表。"
    ))
    _table(
        doc,
        retrieval,
        ["方案", "作用", "前1名召回率", "前3名召回率", "前10名召回率", "前10名平均倒数排名"],
        [3.5, 5.8, 2, 2, 2, 2.5],
    )
    _body(doc, (
        "ColSmol以Recall@1=0.3320、Recall@3=0.3900和MRR@10=0.3669取得最佳结果。InfoNCE部分"
        "微调使SigLIP的Recall@10由0.2520提高到0.3240，说明领域图文配对训练有效；加入一个难负例"
        "后MRR@10为0.2313，没有超过普通部分微调的0.2320，表明当前难负例采样尚未带来额外收益。"
    ))

    _heading(doc, "8.2 检索失败案例的原因分解", 2)
    counts = failure_summary["category_counts"]
    labels = {
        "all_methods_miss": "所有方法前10均失败",
        "colsmol_recovers": "ColSmol救回前三",
        "finetune_recovers": "InfoNCE微调救回前三",
        "ocr_advantage": "OCR文本匹配占优",
        "ranked_below_top3": "证据位于第4至10名",
        "top3_success": "至少一种方案前三命中",
    }
    failure_rows = [
        {
            "失败类型": labels[key],
            "数量": value,
            "占测试集比例": f"{value / failure_summary['test_query_count']:.1%}",
            "主要判断": failure_summary["interpretation"][key],
        }
        for key, value in counts.items()
    ]
    _table(doc, failure_rows, ["失败类型", "数量", "占测试集比例", "主要判断"], [4.2, 1.7, 2.4, 9.7])
    _body(doc, (
        "在500条测试Query中，254条在具有逐条结果的SigLIP零样本、微调SigLIP、OCR+BGE和ColSmol"
        "四条链路中均未进入前10；54条只有ColSmol将证据提升到前三，体现局部视觉词元晚交互的优势；"
        "30条由InfoNCE微调救回前三；3条由OCR精确文本匹配获得更好排名；37条证据位于第4至10名，"
        "说明扩大Top-K可以增加召回，但会同步提高生成开销。难负例实验未保存逐Query结果，因此只进入"
        "总体基线，不参与本节逐条归因。"
    ))

    _heading(doc, "8.3 四种SmolVLM问答方案复核", 2)
    _table(doc, generation, ["方案", "作用", "证据召回率", "完全匹配率", "宽松准确率", "命中证据后的准确率", "平均生成耗时(ms)"], [2.8, 4.3, 2, 2, 2, 2.8, 2.9])
    _body(doc, (
        "Oracle Page宽松准确率为0.3460，说明生成模型本身已构成上限约束。ColSmol Top-1宽松准确率"
        "为0.1840，Top-3逐页融合仅提高至0.1860，却使平均耗时由455.1 ms增至1369.5 ms；Top-3拼图"
        "准确率下降至0.1260。就当前单页实验而言，ColSmol Top-1仍是更合理的速度与准确率折中。"
    ))

    _heading(doc, "8.4 连续三页财务问答验证集", 2)
    _body(doc, (
        f"验证集包含{multipage_manifest['document_count']}份连续三页财务报告、"
        f"{multipage_manifest['page_count']}张页面图像和{multipage_manifest['sample_count']}条跨页问题。"
        "每份报告的第1页给出2025年收入、成本和员工数，第2页给出2026年对应数据，第3页给出研发支出"
        "等补充指标。问题覆盖收入增量、经营利润变化、员工增量，以及经营利润与研发投入求和。每条"
        "Query均标注两个必要证据页，单看任意一页都不足以推出标准答案。"
    ))
    for page_no in (1, 2, 3):
        image_path = MULTIPAGE / "pages" / f"finance_01_aster_p{page_no}.png"
        picture = doc.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(str(image_path), width=Cm(8.4))
        caption = doc.add_paragraph(f"图8-{page_no}  Aster Holdings连续财务报告第{page_no}页")
        caption.style = doc.styles["Caption"]
        _apply_ppr(caption, CAPTION_PPR)
    examples = [
        {
            "问题": "2026年收入比2025年增加多少？",
            "答案": "10",
            "必要证据": "第1页、第2页",
            "推理": "第2页收入123减第1页收入113",
        },
        {
            "问题": "2025至2026年经营利润变化多少？",
            "答案": "4",
            "必要证据": "第1页、第2页",
            "推理": "分别计算两年收入减成本后作差",
        },
        {
            "问题": "2026年经营利润与研发支出之和是多少？",
            "答案": "67",
            "必要证据": "第2页、第3页",
            "推理": "第2页经营利润57加第3页研发10",
        },
    ]
    _table(doc, examples, ["问题", "答案", "必要证据", "推理"], [7.2, 1.5, 3.2, 6.1])

    _heading(doc, "8.5 多页检索与联合问答结果", 2)
    result_rows = [
        {"指标": "任意证据召回率", "结果": f"{metrics['any_evidence_recall']:.4f}", "解释": "Top-3至少包含一个必要证据页"},
        {"指标": "全部证据召回率", "结果": f"{metrics['all_evidence_recall']:.4f}", "解释": "Top-3同时包含两个必要证据页"},
        {"指标": "平均证据覆盖率", "结果": f"{metrics['mean_evidence_coverage']:.4f}", "解释": "必要证据页的平均召回比例"},
        {"指标": "跨页宽松准确率", "结果": f"{metrics['relaxed_accuracy']:.4f}", "解释": "SmolVLM多图回答的数值正确率"},
        {"指标": "平均生成耗时", "结果": f"{metrics['mean_generation_ms']:.1f} ms", "解释": "RTX 4070上每题平均多图推理耗时"},
    ]
    _table(doc, result_rows, ["指标", "结果", "解释"], [5, 3, 10])
    _body(doc, (
        "ColSmol在40题中全部至少召回一个证据页，并在39题中同时找齐全部必要证据页，全部证据召回率"
        "达到0.9750。然而SmolVLM-500M的宽松准确率为0.0000，即使只考察已找齐全部证据的39题，"
        "准确率仍为0。逐条输出显示，模型通常抄写某一页的收入或员工总数，没有执行跨页减法；部分输出"
        "包含算式，但选错操作数或混淆收入、利润和研发支出。40条预测中没有一条仅因携带单位或解释而"
        "被误判，目标数值本身均未正确出现，因此该负结果可信。"
    ))
    type_rows = _type_rows(predictions)
    _table(doc, type_rows, ["问题类型", "样本数", "正确数", "准确率", "典型错误"], [4, 2, 2, 2, 8])

    _heading(doc, "8.6 二阶段结论", 2)
    _body(doc, (
        "本阶段完成了五种检索方案统一比较、检索失败原因归类、四种SmolVLM问答方案复核，以及真正"
        "需要多个证据页的连续文档验证。实验确认ColSmol是当前最强页面检索器，InfoNCE微调对SigLIP"
        "有效，但简单难负例采样未继续提升；在受控多页数据上，系统基本跨过多证据检索门槛，却尚未"
        "跨过联合推理门槛。因此当前项目可以表述为已完成页面级视觉RAG原型和多证据检索验证，但不能"
        "宣称已经解决真实多页财务问答。"
    ))


def _type_rows(rows: list[dict[str, object]]) -> list[dict[str, object]]:
    grouped: dict[int, list[dict[str, object]]] = defaultdict(list)
    for row in rows:
        grouped[int(row["query_id"].rsplit("q", 1)[1])].append(row)
    labels = {
        1: ("收入增量", "抄写单年收入，未执行跨页减法"),
        2: ("经营利润变化", "未分别计算两年利润或错误相减"),
        3: ("员工增量", "输出员工总数而非年度增量"),
        4: ("利润与研发求和", "混淆收入、利润和研发支出"),
    }
    result = []
    for key in sorted(grouped):
        values = grouped[key]
        correct = sum(bool(row["relaxed_correct"]) for row in values)
        result.append({
            "问题类型": labels[key][0],
            "样本数": len(values),
            "正确数": correct,
            "准确率": f"{correct / len(values):.1%}",
            "典型错误": labels[key][1],
        })
    return result


def _heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    _apply_ppr(paragraph, HEADING1_PPR if level == 1 else HEADING2_PPR)
    _format_heading_runs(paragraph, 18 if level == 1 else 15)


def _format_heading_runs(paragraph, size: int) -> None:
    for run in paragraph.runs:
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(size)


def _body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    _apply_ppr(paragraph, BODY_PPR)


def _table(doc: Document, rows: list[dict[str, object]], columns: list[str], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    _replace_child(table._tbl, "w:tblPr", TABLE_PR)
    for index, column in enumerate(columns):
        cell = table.rows[0].cells[index]
        cell.text = column
        cell.width = Cm(widths[index])
        _replace_child(cell._tc, "w:tcPr", TABLE_HEADER_TCPR)
        _apply_ppr(cell.paragraphs[0], TABLE_HEADER_PPR)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(9)
    _replace_child(table.rows[0]._tr, "w:trPr", TABLE_HEADER_TRPR)
    for row in rows:
        cells = table.add_row().cells
        _replace_child(table.rows[-1]._tr, "w:trPr", TABLE_BODY_TRPR)
        for index, column in enumerate(columns):
            value = row.get(column, "")
            cells[index].text = str(value)
            cells[index].width = Cm(widths[index])
            _replace_child(cells[index]._tc, "w:tcPr", TABLE_BODY_TCPR)
            _apply_ppr(cells[index].paragraphs[0], TABLE_BODY_PPR)
            for run in cells[index].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(9)
    spacer = doc.add_paragraph()
    _apply_ppr(spacer, BODY_PPR)


def _three_line(table) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.first_child_found_in("w:tblBorders")
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    tbl_pr.append(borders)
    for edge in ("top", "insideH", "bottom"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "12" if edge in ("top", "bottom") else "4")
        element.set(qn("w:color"), "000000")
        borders.append(element)


def _capture_format_templates(doc: Document) -> None:
    global HEADING1_PPR, HEADING2_PPR, BODY_PPR, CAPTION_PPR
    global TABLE_PR, TABLE_HEADER_TRPR, TABLE_BODY_TRPR
    global TABLE_HEADER_TCPR, TABLE_BODY_TCPR, TABLE_HEADER_PPR, TABLE_BODY_PPR

    heading1 = next(p for p in doc.paragraphs if p.style.name == "Heading 1")
    heading2 = next(p for p in doc.paragraphs if p.style.name == "Heading 2")
    body = next(p for p in doc.paragraphs if p.style.name == "Normal" and len(p.text) > 80)
    caption = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("图") and len(p.text.strip()) > 4
    )
    HEADING1_PPR = deepcopy(heading1._p.pPr)
    HEADING2_PPR = deepcopy(heading2._p.pPr)
    BODY_PPR = deepcopy(body._p.pPr)
    CAPTION_PPR = deepcopy(caption._p.pPr)

    template = next(table for table in doc.tables if len(table.columns) >= 4)
    TABLE_PR = deepcopy(template._tbl.tblPr)
    TABLE_HEADER_TRPR = deepcopy(template.rows[0]._tr.trPr)
    TABLE_BODY_TRPR = deepcopy(template.rows[1]._tr.trPr)
    TABLE_HEADER_TCPR = deepcopy(template.rows[0].cells[0]._tc.tcPr)
    TABLE_BODY_TCPR = deepcopy(template.rows[1].cells[0]._tc.tcPr)
    TABLE_HEADER_PPR = deepcopy(template.rows[0].cells[0].paragraphs[0]._p.pPr)
    TABLE_BODY_PPR = deepcopy(template.rows[1].cells[0].paragraphs[0]._p.pPr)


def _apply_ppr(paragraph, template) -> None:
    if template is None:
        return
    existing = paragraph._p.pPr
    if existing is not None:
        paragraph._p.remove(existing)
    paragraph._p.insert(0, deepcopy(template))


def _replace_child(parent, tag: str, template) -> None:
    if template is None:
        return
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)
    parent.insert(0, deepcopy(template))


def _csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        return list(csv.DictReader(file))


def _json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
